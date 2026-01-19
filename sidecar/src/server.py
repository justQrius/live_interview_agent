"""
WebSocket server for Live Interview Agent sidecar.

Handles bidirectional communication with the Tauri UI application.
Coordinates audio capture, STT, RAG, and LLM processing.
"""

import asyncio
import json
import logging
import base64
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Set, cast
import numpy as np

import websockets
from websockets.asyncio.server import serve, ServerConnection

# Allow running directly from command line
import sys
import os
if __name__ == "__main__" and __package__ is None:
    # Add the sidecar directory (parent of src) to sys.path
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.protocol import (
    Message,
    MessageType,
    SessionStatus,
    ConfidenceLevel,
    Speaker,
    EnhancementType,
    create_transcription_message,
    create_answer_chunk_message,
    create_error_message,
    create_status_message,
    create_session_list_message,
    create_session_data_message,
    create_session_export_message,
    create_session_deleted_message,
    create_preparation_ready_message,
    create_extraction_progress_message,
    create_extraction_complete_message,
    create_interim_transcription_message,
    create_story_suggestion_message,
    create_structure_suggestion_message,
    create_consistency_warning_message,
    create_enhanced_answer_start_message,
    create_enhanced_answer_chunk_message,
    create_enhanced_answer_complete_message,
    create_document_type_suggestions_message,
)
from src.audio.diarization import SpeakerRecognizer
from src.audio.capture import AudioCapture, AudioCaptureError
from src.audio.vad import VADProcessor, SpeechSegment
from src.audio.noise_reduction import NoiseReducer
from src.providers.base import STTProvider, LLMProvider
from src.providers.factory import ProviderFactory
from src.providers.config import ProviderConfig, GeminiModels
from src.context.enhanced_manager import EnhancedContextManager, DocumentType as ContextDocumentType
from src.context.manager import ContextManager
from src.rag.store import VectorStore
from src.rag.engine import RAGEngine
from src.rag.enhanced_engine import EnhancedRAGEngine
from src.providers.stt.gemini import GeminiSTTProvider
from src.rag.speculative import SpeculativeRetriever
from src.warmup import ModelWarmer
from src.classification.question_detector import QuestionDetector
from src.classification.query_reformulator import QueryReformulator
from src.classification.question_splitter import QuestionSplitter
from src.providers.llm.prompts import classify_question
from src.storage.session_store import SessionHistoryStore
from src.storage.exporter import SessionExporter, ExportFormat
from src.memory.store import MemoryStore
from src.memory.models import DocumentType as MemoryDocumentType
from src.extraction.pipeline import ExtractionPipeline
from src.coaching.story_recaller import StoryRecaller
from src.coaching.structure_suggester import StructureSuggester
from src.coaching.consistency_tracker import ConsistencyTracker
from src.classification.document_classifier import DocumentClassifier

# Phase 7: Streaming STT for low-latency semantic endpointing
from src.providers.stt.streaming_manager import StreamingSTTManager, StreamingSTTCallbacks
from src.providers.stt.streaming_base import StreamingConfig, EndOfTurnEvent, EndpointingType

# Phase 6: Utterance Accumulation
from src.classification.utterance_accumulator import UtteranceAccumulator
from src.classification.accumulator_models import AccumulatorConfig
from src.protocol import create_accumulating_message

# Phase 5: Gemini Features
from src.context.gemini_cache import GeminiCacheManager
from src.context.file_uploader import GeminiFileUploader, DocumentType as FileDocumentType
from src.rag.gemini_embeddings import GeminiEmbeddingFunction
# Note: EnhancedContextManager already imported above (line 60)

# Phase 8: RAG Persistence
from src.storage.rag_manifest import RagManifest
from src.protocol import (
    create_rag_state_message,
    create_cache_refresh_complete_message,
    create_data_cleared_message,
)

# Configure logging with file output for debugging
import logging.handlers
from pathlib import Path

# Create logs directory in sidecar folder
LOG_DIR = Path(__file__).parent.parent / "logs"
LOG_DIR.mkdir(exist_ok=True)
LOG_FILE = LOG_DIR / "sidecar.log"

# Configure root logger with both console and file handlers
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[
        # Console handler (existing behavior)
        logging.StreamHandler(),
        # Rotating file handler: 5MB per file, keep last 5 files
        logging.handlers.RotatingFileHandler(
            LOG_FILE,
            maxBytes=5 * 1024 * 1024,  # 5 MB
            backupCount=5,
            encoding="utf-8"
        )
    ]
)
logger = logging.getLogger(__name__)
logger.info(f"Logging to file: {LOG_FILE}")


@dataclass
class SessionState:
    """Tracks the current session state."""

    status: SessionStatus = SessionStatus.IDLE
    api_key: Optional[str] = None
    voice_calibrated: bool = False
    user_embedding: Optional[np.ndarray] = field(default=None, repr=False)
    # Conversation history for LLM context (list of {"role": "user"|"assistant", "content": str})
    conversation_history: List[Dict[str, str]] = field(default_factory=list)
    # Phase 3: Persistent session ID for history storage
    persistent_session_id: Optional[str] = None
    # Phase 3B: Pre-interview preparation summary
    preparation_summary: Optional[str] = None
    # Listening control: when True, audio/STT is paused but manual input works
    listening_paused: bool = False


class SidecarServer:
    """
    WebSocket server for the Live Interview Agent sidecar.

    Handles client connections and routes messages to appropriate handlers.
    """

    def __init__(self, host: str = "127.0.0.1", port: int = 8765):
        """
        Initialize the sidecar server.

        Args:
            host: Host to bind to (default: localhost only for security)
            port: Port to listen on (default: 8765)
        """
        self.host = host
        self.port = port
        self._audio_task: Optional[asyncio.Task] = None
        self._background_tasks: Set[asyncio.Task] = set()
        self.clients: Set[ServerConnection] = set()
        self.session_state = SessionState()
        self._server: Optional[Any] = None
        self._running = False
        
        self.provider_factory: Optional[ProviderFactory] = None
        self.stt: Optional[STTProvider] = None
        self.vad: Optional[VADProcessor] = None
        self.noise_reducer: Optional[NoiseReducer] = None
        self.audio_capture: Optional[AudioCapture] = None
        self.llm: Optional[LLMProvider] = None
        self.vector_store: Optional[VectorStore] = None
        self.rag_engine: Optional[EnhancedRAGEngine] = None
        self.speculative_retriever: Optional[SpeculativeRetriever] = None
        self.story_recaller: Optional[StoryRecaller] = None
        self.structure_suggester: Optional[StructureSuggester] = None
        
        # Phase 5: Gemini Managers
        self.gemini_cache_manager: Optional[GeminiCacheManager] = None
        self.gemini_file_uploader: Optional[GeminiFileUploader] = None
        
        # Initialize storage and managers
        self.session_store = SessionHistoryStore()
        self.context_manager = EnhancedContextManager()
        
        # Phase 4E: Initialize Consistency Tracker (needs session store)
        self.consistency_tracker = ConsistencyTracker(self.session_store)
        
        self.model_warmer = ModelWarmer.get_instance()
        self.model_warmer.start_warming()
        
        self.speaker_recognizer = None 
        
        # Phase 3: Intelligent Question Detection
        self.question_detector = QuestionDetector()
        self.question_detection_enabled = True  # Feature flag for rollout
        self.question_confidence_threshold = 0.7  # Configurable threshold
        
        # Phase 3C: Conversational Intelligence
        self.query_reformulator = QueryReformulator()
        self.question_splitter = QuestionSplitter()
        
        # Phase 6: Utterance Accumulation for multi-segment questions
        self.utterance_accumulator = UtteranceAccumulator()
        self.accumulation_enabled = True  # Feature flag for rollout
        
        self.session_persistence_enabled = True  # Feature flag for rollout 
        
        # Phase 4: Persistent Memory & Extraction Pipeline
        self.memory_store = MemoryStore()
        self.extraction_pipeline = ExtractionPipeline(memory_store=self.memory_store)
        self.extraction_enabled = True  # Feature flag for rollout
        
        # Phase 5: Document Type Classifier (initialized without LLM, set later)
        self.document_classifier = DocumentClassifier()
        
        # Phase 5: Enhancement task tracking for cancellation
        self._enhancement_task: Optional[asyncio.Task] = None
        
        # Phase 7: Streaming STT Manager for low-latency semantic endpointing
        self.streaming_stt_manager: Optional[StreamingSTTManager] = None
        self.streaming_stt_enabled = True  # Feature flag for rollout
        
        # Phase 8: RAG Persistence
        self.rag_manifest = RagManifest()

    def _create_background_task(self, coro) -> asyncio.Task:
        """Create and track a background task."""
        task = asyncio.create_task(coro)
        self._background_tasks.add(task)
        task.add_done_callback(self._background_tasks.discard)
        return task

    async def start(self) -> None:
        """Start the WebSocket server."""
        self._running = True
        self._server = await serve(
            self._handle_client,
            self.host,
            self.port,
            max_size=20 * 1024 * 1024  # 20MB limit for large context uploads
        )
        logger.info(f"Sidecar server started on ws://{self.host}:{self.port}")

        while self._running:
            await asyncio.sleep(0.1)

    async def stop(self) -> None:
        """Stop the WebSocket server."""
        self._running = False

        # Stop audio processing first
        await self._stop_audio_processing()

        # Cancel all background tasks
        if self._background_tasks:
            # Create a copy of the set to avoid "Set size changed during iteration"
            tasks = list(self._background_tasks)
            for task in tasks:
                if not task.done():
                    task.cancel()
            
            if tasks:
                await asyncio.gather(*tasks, return_exceptions=True)
            self._background_tasks.clear()

        if self.clients:
            await asyncio.gather(
                *[client.close() for client in self.clients],
                return_exceptions=True
            )
            self.clients.clear()

        if self._server:
            self._server.close()
            self._server = None

        self.session_state = SessionState()
        self.context_manager.clear_context()

        logger.info("Sidecar server stopped")

    async def _handle_client(self, websocket: ServerConnection) -> None:
        """
        Handle a client connection.

        Args:
            websocket: The WebSocket connection
        """
        self.clients.add(websocket)
        logger.info(f"Client connected: {websocket.remote_address}")

        try:
            async for raw_message in websocket:
                if isinstance(raw_message, bytes):
                    raw_message = raw_message.decode('utf-8')
                await self._process_message(websocket, raw_message)
        except websockets.exceptions.ConnectionClosed:
            logger.info(f"Client disconnected: {websocket.remote_address}")
        except Exception as e:
            logger.error(f"Error handling client: {e}")
        finally:
            self.clients.discard(websocket)

    async def _process_message(
        self,
        websocket: ServerConnection,
        raw_message: str
    ) -> None:
        """
        Process an incoming message.

        Args:
            websocket: The WebSocket connection
            raw_message: The raw message string
        """
        try:
            message = Message.from_json(raw_message)
        except json.JSONDecodeError as e:
            logger.warning(f"Invalid JSON received: {e}")
            error_msg = create_error_message(
                f"Invalid JSON: {e}",
                code="ERR_INVALID_JSON"
            )
            await websocket.send(error_msg.to_json())
            return
        except ValueError as e:
            logger.warning(f"Unknown message type: {e}")
            error_msg = create_error_message(
                f"Unknown message type: {e}",
                code="ERR_UNKNOWN_TYPE"
            )
            await websocket.send(error_msg.to_json())
            return

        handlers = {
            MessageType.START_SESSION: self._handle_start_session,
            MessageType.STOP_SESSION: self._handle_stop_session,
            MessageType.UPLOAD_CONTEXT: self._handle_upload_context,
            MessageType.CALIBRATE_VOICE: self._handle_calibrate_voice,
            MessageType.MANUAL_QUESTION: self._handle_manual_question,
            # Session History handlers (Phase 3: STORY-039)
            MessageType.LIST_SESSIONS: self._handle_list_sessions,
            MessageType.LOAD_SESSION: self._handle_load_session,
            MessageType.EXPORT_SESSION: self._handle_export_session,
            MessageType.DELETE_SESSION: self._handle_delete_session,
            # Pre-interview preparation (Phase 3B: STORY-047)
            MessageType.PREPARE_INTERVIEW: self._handle_prepare_interview,
            # Phase 5: Answer Enhancement
            MessageType.ENHANCE_ANSWER: self._handle_enhance_answer,
            MessageType.CANCEL_ENHANCEMENT: self._handle_cancel_enhancement,
            # Phase 5: Document Type Inference
            MessageType.INFER_DOCUMENT_TYPES: self._handle_infer_document_types,
            # Phase 8: RAG Persistence
            MessageType.LOAD_RAG_STATE: self._handle_load_rag_state,
            MessageType.REFRESH_CACHE: self._handle_refresh_cache,
            MessageType.CLEAR_ALL_DATA: self._handle_clear_all_data,
            # Listening Control
            MessageType.PAUSE_LISTENING: self._handle_pause_listening,
            MessageType.RESUME_LISTENING: self._handle_resume_listening,
        }

        handler = handlers.get(message.type)
        if handler:
            await handler(websocket, message)
        else:
            error_msg = create_error_message(
                f"Unsupported message type: {message.type}",
                code="ERR_UNSUPPORTED"
            )
            await websocket.send(error_msg.to_json())

    async def _handle_start_session(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle START_SESSION message."""
        data = message.data or {}
        
        if "apiKeys" not in data and "apiKey" in data:
            data = {
                "apiKeys": {"gemini": data["apiKey"]},
                "preferences": {}
            }
            
        try:
            config = ProviderConfig.from_dict(data)
            self.provider_factory = ProviderFactory(config)
            
            self.stt = self.provider_factory.get_stt_provider()
            
            # Phase 5: Initialize Gemini Managers if key available
            # Reuse existing managers if API key hasn't changed to preserve context
            if config.gemini_api_key:
                try:
                    # Check if we can reuse existing cache manager
                    should_init_cache = True
                    if self.gemini_cache_manager:
                        # If manager exists and key matches (or we assume same key for simplicity in session reuse)
                        # In strict mode we should check key, but for now reuse is priority
                        should_init_cache = False
                        logger.info("Reusing existing Gemini Cache Manager")
                    
                    if should_init_cache:
                        self.gemini_cache_manager = GeminiCacheManager(config.gemini_api_key)
                    
                    # Check if we can reuse existing file uploader
                    should_init_uploader = True
                    if self.gemini_file_uploader:
                        should_init_uploader = False
                        logger.info("Reusing existing Gemini File Uploader")
                        
                    if should_init_uploader:
                        self.gemini_file_uploader = GeminiFileUploader(config.gemini_api_key)
                        
                    if should_init_cache or should_init_uploader:
                        logger.info("Initialized Gemini Cache & Upload managers")
                except Exception as e:
                    logger.warning(f"Failed to init Gemini managers: {e}")
            
            try:
                self.llm = self.provider_factory.get_llm_provider()
                
                # Phase 5: Attach Gemini Cache if available
                # This ensures the LLM has access to the full document context
                if self.gemini_cache_manager and self.gemini_cache_manager.current_cache_name:
                    if hasattr(self.llm, 'set_cached_content'):
                        getattr(self.llm, 'set_cached_content')(self.gemini_cache_manager.current_cache_name)
                        logger.info(f"Attached existing Gemini cache to session LLM: {self.gemini_cache_manager.current_cache_name}")
                
                # Phase 4: Set LLM for Extraction Pipeline
                if self.extraction_pipeline:
                    self.extraction_pipeline.set_llm_provider(self.llm)
                    logger.info("Set LLM provider for Extraction Pipeline")

                # Phase 4: Load existing profile for context injection
                if self.memory_store:
                    existing_profile = self.memory_store.get_profile()
                    if existing_profile:
                        logger.info(f"Loaded existing candidate profile ({len(existing_profile.profile_text)} chars)")
                        self.llm.set_candidate_profile(existing_profile.get_prompt_injection())
                
                # Phase 4C: Set LLM for Tier 3 Question Detection
                self.question_detector.set_llm_provider(self.llm)
                
                # Phase 6: Set LLM for Tier 4 Completeness Detection (Utterance Accumulator)
                if hasattr(self, 'utterance_accumulator') and self.utterance_accumulator:
                    self.utterance_accumulator.set_llm_provider(self.llm)
                
                # Phase 5 Fix: Create Gemini cache if context was uploaded BEFORE session start
                # This fixes the timing bug where UPLOAD_CONTEXT happens before START_SESSION
                if (self.gemini_cache_manager and 
                    self.gemini_file_uploader and 
                    self.gemini_file_uploader.has_files() and
                    not self.gemini_cache_manager.current_cache_name and
                    hasattr(self.llm, 'set_cached_content')):
                    logger.info("Context was uploaded before session start - creating Gemini cache now")
                    self._create_background_task(self._create_gemini_cache_background())
                
            except Exception as e:
                logger.warning(f"No LLM provider available: {e}")
                self.llm = None

        except Exception as e:
            logger.error(f"Failed to initialize providers: {e}")
            error_msg = create_error_message(
                f"Failed to initialize providers: {e}",
                code="ERR_PROVIDER_INIT"
            )
            await websocket.send(error_msg.to_json())
            return

        self.session_state.api_key = "multi-provider-active" 
        self.session_state.status = SessionStatus.LISTENING
        
        # CRITICAL: Start audio processing FIRST - this is the user's primary need
        # RAG initialization can happen in background and may timeout on network issues
        try:
            await self._start_audio_processing()
        except Exception as e:
            logger.error(f"Failed to start audio processing: {e}")
            error_msg = create_error_message(
                f"Failed to start audio processing: {e}",
                code="ERR_AUDIO_START"
            )
            await websocket.send(error_msg.to_json())
            self.session_state.status = SessionStatus.IDLE
            return

        logger.info("Session started - audio processing active")
        
        # Phase 7: Initialize Streaming STT for low-latency semantic endpointing
        if self.streaming_stt_enabled and self.provider_factory:
            try:
                self.streaming_stt_manager = StreamingSTTManager(self.provider_factory)
                
                # Create callbacks for streaming events
                callbacks = StreamingSTTCallbacks(
                    on_interim=self._on_streaming_interim,
                    on_final=self._on_streaming_final,
                    on_end_of_turn=self._on_streaming_end_of_turn,
                    on_error=self._on_streaming_error,
                )
                
                # Start streaming session in background (non-blocking)
                streaming_started = await self.streaming_stt_manager.start_session(callbacks)
                if streaming_started:
                    logger.info(f"Streaming STT started: {self.streaming_stt_manager.provider_name} "
                               f"(semantic={self.streaming_stt_manager.supports_semantic_endpointing})")
                    
                    # Phase 7: Auto-adjust accumulator based on provider capability
                    # Semantic providers (AssemblyAI, OpenAI) handle endpointing well → disable accumulator
                    # Acoustic providers (Deepgram) only detect pauses → keep accumulator for safety
                    if self.streaming_stt_manager.supports_semantic_endpointing:
                        if hasattr(self, 'utterance_accumulator') and self.utterance_accumulator:
                            self.utterance_accumulator.config.endpointing_mode = "streaming"
                            logger.info("Auto-switched to streaming endpointing mode (semantic provider detected)")
                    else:
                        if hasattr(self, 'utterance_accumulator') and self.utterance_accumulator:
                            self.utterance_accumulator.config.endpointing_mode = "hybrid"
                            logger.info("Using hybrid endpointing mode (acoustic provider detected)")
                else:
                    logger.info("Streaming STT not available, using batch STT only")
            except Exception as e:
                logger.warning(f"Failed to start streaming STT: {e}")
                self.streaming_stt_manager = None

        status_msg = create_status_message(self.session_state.status)
        await websocket.send(status_msg.to_json())
        
        # Phase 3: Create persistent session for history storage
        if self.session_persistence_enabled:
            try:
                context_files = [f.get("name", "") for f in data.get("files", [])]
                self.session_state.persistent_session_id = self.session_store.create_session(
                    context_files=context_files
                )
                logger.info(f"Created persistent session: {self.session_state.persistent_session_id}")
                
                # Phase 4E: Start consistency tracking
                if self.consistency_tracker:
                    self.consistency_tracker.start_session(self.session_state.persistent_session_id)
                    
            except Exception as e:
                logger.warning(f"Failed to create persistent session: {e}")
                # Continue without persistence - don't break main flow
        
        # Initialize RAG in background - don't block audio processing
        # Session works even if RAG fails (graceful degradation)
        # Reuse existing RAG engine if available to preserve context
        if not self.rag_engine or not self.vector_store:
            rag_key = config.gemini_api_key or config.openai_api_key or "dummy"
            self._create_background_task(self._init_rag_background(rag_key))
        else:
            logger.info("RAG engine already active, skipping initialization")

    async def _handle_stop_session(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle STOP_SESSION message."""
        # Phase 3: End persistent session before clearing state
        if self.session_persistence_enabled and self.session_state.persistent_session_id:
            try:
                self.session_store.end_session(self.session_state.persistent_session_id)
                logger.info(f"Ended persistent session: {self.session_state.persistent_session_id}")
            except Exception as e:
                logger.warning(f"Failed to end persistent session: {e}")
        
        self.session_state.status = SessionStatus.IDLE
        self.session_state.api_key = None
        # Clear conversation history for fresh start on next session
        self.session_state.conversation_history.clear()
        self.session_state.persistent_session_id = None
        
        # NOTE: We preserve context (RAG, Cache, Profile) across session stops 
        # to allow restarting without re-uploading everything.
        # Cleanup happens only on process exit.
        
        # Only clear ephemeral conversation history
        self.session_state.conversation_history.clear()
        
        # Don't clear Context/Providers here
        # self.context_manager.clear_context()
        # if self.gemini_file_uploader: ...
        # if self.gemini_cache_manager: ...
        # if self.vector_store: ...
        # if self.llm: self.llm.clear_candidate_profile()
        
        # Stop audio processing
        await self._stop_audio_processing()
        
        # Phase 7: Stop streaming STT session
        if self.streaming_stt_manager:
            try:
                await self.streaming_stt_manager.stop_session()
                logger.info("Stopped streaming STT session")
            except Exception as e:
                logger.warning(f"Error stopping streaming STT: {e}")
        
        # Phase 6: Reset utterance accumulator
        if hasattr(self, 'utterance_accumulator') and self.utterance_accumulator:
            self.utterance_accumulator.reset()
        
        # Cancel any ongoing enhancement task
        if self._enhancement_task and not self._enhancement_task.done():
            self._enhancement_task.cancel()
            try:
                await self._enhancement_task
            except asyncio.CancelledError:
                pass
            self._enhancement_task = None
            logger.info("Cancelled ongoing enhancement due to session stop")

        logger.info("Session stopped (Context preserved for restart)")

        status_msg = create_status_message(SessionStatus.IDLE)
        await websocket.send(status_msg.to_json())

    async def _handle_upload_context(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle UPLOAD_CONTEXT message."""
        data = message.data or {}
        files = data.get("files", [])

        # Feature: Auto-initialize providers if apiKeys present in upload message
        if "apiKeys" in data and not self.llm:
            try:
                logger.info("Found API keys in upload message, initializing providers...")
                config = ProviderConfig.from_dict(data)
                self.provider_factory = ProviderFactory(config)
                self.llm = self.provider_factory.get_llm_provider()
                if self.extraction_pipeline:
                    self.extraction_pipeline.set_llm_provider(self.llm)
                
                # Also init STT if needed (though less critical for upload)
                if not self.stt:
                    self.stt = self.provider_factory.get_stt_provider()
                
                # Init Gemini managers
                # FIX: Always initialize if not already present, regardless of cache manager state
                if config.gemini_api_key:
                    if not self.gemini_cache_manager:
                        self.gemini_cache_manager = GeminiCacheManager(config.gemini_api_key)
                        logger.info("Created Gemini Cache Manager for upload")
                    if not self.gemini_file_uploader:
                        self.gemini_file_uploader = GeminiFileUploader(config.gemini_api_key)
                        logger.info("Created Gemini File Uploader for upload")
                
                logger.info("Providers initialized from upload message")
            except Exception as e:
                logger.warning(f"Failed to auto-init providers from upload: {e}")

        if not files:
            logger.warning("No files provided in context upload")
            error_msg = create_error_message(
                "No files provided",
                code="ERR_NO_FILES"
            )
            await websocket.send(error_msg.to_json())
            return

        logger.info(f"Context upload requested: {len(files)} files")
        
        self.session_state.status = SessionStatus.PROCESSING
        await websocket.send(create_status_message(SessionStatus.PROCESSING).to_json())

        processed_count = 0
        total_chunks = 0
        errors = []
        
        # Type mapping dicts (defined once, used for all files)
        context_doc_type_map = {
            "resume": ContextDocumentType.RESUME,
            "job_description": ContextDocumentType.JOB_DESCRIPTION,
            "company_info": ContextDocumentType.COMPANY_INFO,
            "interviewer_info": ContextDocumentType.INTERVIEWER_INFO,
            "sample_qa": ContextDocumentType.SAMPLE_QA,
        }
        memory_doc_type_map = {
            "resume": MemoryDocumentType.RESUME,
            "job_description": MemoryDocumentType.JOB_DESCRIPTION,
            "company_info": MemoryDocumentType.COMPANY_INFO,
            "interviewer_info": MemoryDocumentType.INTERVIEWER_INFO,
        }
        file_doc_type_map = {
            "resume": FileDocumentType.RESUME,
            "job_description": FileDocumentType.JOB_DESCRIPTION,
            "company_info": FileDocumentType.COMPANY_INFO,
            "interviewer_info": FileDocumentType.INTERVIEWER_INFO,
            "sample_qa": FileDocumentType.SAMPLE_QA,
        }
        
        # Collect Gemini upload tasks for parallel execution
        # Each entry: (filename, file_doc_type, upload_coroutine)
        gemini_upload_tasks: List[tuple] = []
        # Track file info for post-upload processing: (filename, content, memory_doc_type)
        files_for_extraction: List[tuple] = []

        try:
            # ============================================================
            # PHASE 1: Pre-process files (validation, chunking, indexing)
            # This is fast and can stay sequential
            # ============================================================
            for file_data in files:
                filename = file_data.get("name")
                content = file_data.get("content")
                doc_type_str = file_data.get("documentType", "other")
                
                if not filename or not content:
                    logger.warning(f"Invalid file data: {filename}")
                    errors.append(f"Invalid data for {filename or 'unknown file'}")
                    continue
                
                # Map document types
                context_doc_type = context_doc_type_map.get(doc_type_str, ContextDocumentType.CUSTOM)
                memory_doc_type = memory_doc_type_map.get(doc_type_str, MemoryDocumentType.OTHER)
                file_doc_type = file_doc_type_map.get(doc_type_str, FileDocumentType.CUSTOM)
                    
                try:
                    # Process through context manager for RAG chunks
                    new_chunks = await self.context_manager.process_file(filename, content, document_type=context_doc_type)
                    
                    if new_chunks and self.vector_store:
                        chunk_texts = [c.text for c in new_chunks]
                        chunk_metas = [c.metadata for c in new_chunks]
                        try:
                            # Run blocking vector store add in thread pool
                            loop = asyncio.get_running_loop()
                            store = self.vector_store
                            await loop.run_in_executor(
                                None,
                                lambda s=store, ct=chunk_texts, cm=chunk_metas: s.add_documents(ct, metadatas=cm)
                            )
                            logger.info(f"Indexed {len(new_chunks)} chunks for {filename} (non-blocking)")
                        except Exception as e:
                            logger.error(f"Failed to add chunks to vector store: {e}")
                    
                    total_chunks += len(new_chunks)
                    
                    # ============================================================
                    # PHASE 2: Queue Gemini upload tasks (DON'T await yet)
                    # ============================================================
                    if self.gemini_file_uploader:
                        # Create the coroutine but DON'T await - we'll batch them
                        upload_coro = self.gemini_file_uploader.upload_from_base64_async(
                            content_b64=content,
                            filename=filename,
                            document_type=file_doc_type
                        )
                        gemini_upload_tasks.append((filename, file_doc_type, upload_coro))
                    
                    # Track for extraction (after all uploads complete)
                    files_for_extraction.append((filename, content, memory_doc_type))
                    
                    # ============================================================
                    # PHASE 8: Write to RAG Manifest for persistence
                    # Store document info so we can restore state on restart
                    # ============================================================
                    try:
                        # Decode content for preview (first 200 chars)
                        preview = ""
                        try:
                            decoded = base64.b64decode(content)
                            # Try to get text preview
                            ext = os.path.splitext(filename)[1].lower()
                            if ext in ['.txt', '.md', '.html', '.json']:
                                preview = decoded.decode('utf-8', errors='ignore')[:200]
                            elif ext == '.docx':
                                import io
                                import docx
                                doc = docx.Document(io.BytesIO(decoded))
                                preview = '\n'.join([p.text for p in doc.paragraphs[:3]])[:200]
                        except Exception:
                            pass
                        
                        self.rag_manifest.add_document(
                            filename=filename,
                            document_type=doc_type_str,
                            file_size_bytes=len(content) if content else 0,
                            chunk_count=len(new_chunks),
                            content_b64=content,  # Store for cache refresh
                            preview=preview,
                        )
                        logger.info(f"Added {filename} to RAG manifest")
                    except Exception as e:
                        logger.warning(f"Failed to add {filename} to manifest: {e}")
                    
                    processed_count += 1
                except Exception as e:
                    logger.error(f"Failed to process {filename}: {e}")
                    errors.append(f"Failed to process {filename}: {str(e)}")
            
            # ============================================================
            # PHASE 3: Execute ALL Gemini uploads in PARALLEL
            # This is the critical performance optimization
            # 18 files * 1.2s each = 22s sequential vs ~2-3s parallel
            # ============================================================
            if gemini_upload_tasks:
                logger.info(f"Starting parallel Gemini upload for {len(gemini_upload_tasks)} files...")
                upload_start = asyncio.get_event_loop().time()
                
                # Extract just the coroutines for gather
                upload_coros = [coro for _, _, coro in gemini_upload_tasks]
                
                # Execute all uploads concurrently, catching individual failures
                upload_results = await asyncio.gather(*upload_coros, return_exceptions=True)
                
                upload_duration = asyncio.get_event_loop().time() - upload_start
                logger.info(f"Parallel Gemini upload completed in {upload_duration:.2f}s")
                
                # Process results
                success_count = 0
                for i, result in enumerate(upload_results):
                    filename, file_doc_type, _ = gemini_upload_tasks[i]
                    if isinstance(result, Exception):
                        logger.warning(f"Failed to upload {filename} to Gemini: {result}")
                        # Don't add to errors list - Gemini upload failure is non-fatal
                    else:
                        logger.info(f"Uploaded {filename} to Gemini File API ({file_doc_type.value})")
                        success_count += 1
                
                logger.info(f"Gemini uploads: {success_count}/{len(gemini_upload_tasks)} successful")
            
            # ============================================================
            # PHASE 4: Start extraction pipelines in background
            # Rate-limited: Process sequentially to avoid API rate limits
            # Previous behavior spawned all concurrently → 429 cascade
            # ============================================================
            if self.extraction_enabled and self.extraction_pipeline and self.llm:
                import uuid
                self.extraction_pipeline.set_llm_provider(self.llm)
                
                # Run extractions sequentially in a single background task
                # to avoid overwhelming the LLM API with parallel requests
                self._create_background_task(self._run_extractions_sequentially(
                    websocket=websocket,
                    files_for_extraction=files_for_extraction,
                ))
            elif self.extraction_enabled and not self.llm:
                logger.warning("Skipping extraction - No LLM provider available yet")

            if processed_count == 0 and errors:
                 error_msg = create_error_message(
                    f"Failed to process files: {'; '.join(errors)}",
                    code="ERR_CONTEXT_PROCESSING"
                )
                 await websocket.send(error_msg.to_json())
            else:
                logger.info(f"Context processed: {processed_count}/{len(files)} files, {total_chunks} chunks")
                if errors:
                    logger.warning(f"Some files failed: {errors}")
                
                # Phase 5 Cache-First: Create Gemini Cache from uploaded files
                # This provides FULL document context with proper attribution
                logger.info(f"Cache check: manager={bool(self.gemini_cache_manager)}, uploader={bool(self.gemini_file_uploader)}, llm={bool(self.llm)}")
                if self.gemini_file_uploader:
                    logger.info(f"Uploader has_files: {self.gemini_file_uploader.has_files()}")
                if self.gemini_cache_manager and self.gemini_file_uploader and self.llm:
                    # Run cache creation in background to avoid blocking UI response
                    self._create_background_task(self._create_gemini_cache_background())
                    logger.info("Triggered Gemini cache creation in background")
                elif self.gemini_cache_manager and self.gemini_file_uploader and not self.llm:
                    # Log why cache creation is deferred - will be created on START_SESSION
                    logger.info("Gemini cache creation deferred - LLM not yet initialized (will create on START_SESSION)")
                else:
                    logger.warning(f"Cannot create cache: manager={bool(self.gemini_cache_manager)}, uploader={bool(self.gemini_file_uploader)}, llm={bool(self.llm)}")
        
        except Exception as e:
            logger.error(f"Context upload fatal error: {e}")
            error_msg = create_error_message(
                f"Context upload failed: {e}",
                code="ERR_CONTEXT_FATAL"
            )
            await websocket.send(error_msg.to_json())

        finally:
            if self.session_state.api_key and self._audio_task:
                 self.session_state.status = SessionStatus.LISTENING
            else:
                 self.session_state.status = SessionStatus.IDLE
                 
            await websocket.send(create_status_message(self.session_state.status).to_json())
    
    async def _run_extractions_sequentially(
        self,
        websocket: ServerConnection,
        files_for_extraction: List[tuple],
    ) -> None:
        """
        Run extraction pipelines sequentially to avoid API rate limits.
        
        Previously, all extractions were launched concurrently, causing 429 errors
        when many files are uploaded at once. This method processes files one at a time
        with a small delay between them.
        
        Args:
            websocket: WebSocket connection for sending progress
            files_for_extraction: List of (filename, content, memory_doc_type) tuples
        """
        import uuid
        
        logger.info(f"Starting sequential extraction for {len(files_for_extraction)} files...")
        
        for i, (filename, content, memory_doc_type) in enumerate(files_for_extraction):
            doc_id = str(uuid.uuid4())
            
            # Create progress callback for this document
            async def make_extraction_progress(ws: ServerConnection, d_id: str, fname: str):
                async def extraction_progress(stage: str, progress: float, msg: str = ""):
                    try:
                        progress_msg = create_extraction_progress_message(
                            stage=stage,
                            progress=progress,
                            document_id=d_id,
                            filename=fname,
                            message=msg
                        )
                        await ws.send(progress_msg.to_json())
                    except Exception as e:
                        logger.debug(f"Failed to send progress: {e}")
                return extraction_progress
            
            extraction_progress = await make_extraction_progress(websocket, doc_id, filename)
            
            logger.info(f"Extracting [{i+1}/{len(files_for_extraction)}]: {filename}")
            
            # Run extraction for this document (awaited, not background)
            await self._run_extraction(
                websocket, doc_id, content, memory_doc_type, filename, extraction_progress
            )
            
            # Small delay between documents to avoid rate limit bursts
            # Skip delay on last document
            if i < len(files_for_extraction) - 1:
                await asyncio.sleep(0.5)
        
        logger.info(f"Sequential extraction complete for {len(files_for_extraction)} files")
    
    async def _run_extraction(
        self,
        websocket: ServerConnection,
        doc_id: str,
        content: str,
        doc_type: MemoryDocumentType,
        filename: str,
        progress_callback
    ) -> None:
        """Run extraction pipeline in background."""
        try:
            result = await self.extraction_pipeline.process_document(
                document_id=doc_id,
                text=content,
                document_type=doc_type,
                filename=filename,
                progress_callback=progress_callback,
            )
            
            # Phase 4: Inject updated profile into LLM if available
            if result.profile and self.llm:
                logger.info(f"Injecting updated candidate profile ({len(result.profile.profile_text)} chars)")
                profile_injection = result.profile.get_prompt_injection()
                self.llm.set_candidate_profile(profile_injection)
                
                # Phase 5 Cache-First: Refresh Gemini cache with updated profile
                if self.gemini_cache_manager:
                    if self.gemini_cache_manager.needs_refresh(profile_injection):
                        logger.info("Profile changed, refreshing Gemini cache...")
                        try:
                            # Prefer file-based cache if files are available
                            if self.gemini_file_uploader and self.gemini_file_uploader.has_files():
                                uploaded_files = self.gemini_file_uploader.get_uploaded_files()
                                document_manifest = self.gemini_file_uploader.get_document_manifest()
                                await self.gemini_cache_manager.create_cache_from_files_async(
                                    uploaded_files=uploaded_files,
                                    document_manifest=document_manifest,
                                    profile_text=profile_injection
                                )
                                logger.info("Gemini file-based cache refreshed with updated profile")
                            elif self.context_manager:
                                # Fallback to chunk-based cache
                                await self.gemini_cache_manager.create_cache_from_context_async(
                                    context_manager=self.context_manager,
                                    profile_text=profile_injection
                                )
                                logger.info("Gemini chunk-based cache refreshed with updated profile")
                        except Exception as cache_err:
                            logger.warning(f"Failed to refresh Gemini cache: {cache_err}")
            
            complete_msg = create_extraction_complete_message(
                document_id=doc_id,
                filename=filename,
                success=result.success,
                summary=result.to_dict()
            )
            await websocket.send(complete_msg.to_json())
            
        except Exception as e:
            logger.error(f"Extraction pipeline failed for {filename}: {e}")
            complete_msg = create_extraction_complete_message(
                document_id=doc_id,
                filename=filename,
                success=False,
                summary={"error": str(e)}
            )
            try:
                await websocket.send(complete_msg.to_json())
            except Exception:
                pass

    async def _create_gemini_cache_background(self) -> None:
        """Background task to create Gemini cache."""
        logger.info("Starting _create_gemini_cache_background task...")
        
        if not (self.gemini_cache_manager and self.gemini_file_uploader and self.llm):
            logger.warning(f"Cache background: missing components - manager={bool(self.gemini_cache_manager)}, uploader={bool(self.gemini_file_uploader)}, llm={bool(self.llm)}")
            return

        if not hasattr(self.llm, 'set_cached_content'):
            logger.warning("Cache background: LLM does not support caching (no set_cached_content method)")
            return
        
        logger.info(f"Cache background: has_files={self.gemini_file_uploader.has_files()}, files_count={len(self.gemini_file_uploader.get_uploaded_files()) if self.gemini_file_uploader.has_files() else 0}")

        if self.gemini_file_uploader.has_files():
            try:
                logger.info("Creating Gemini Cache from uploaded files (Cache-First)...")
                uploaded_files = self.gemini_file_uploader.get_uploaded_files()
                document_manifest = self.gemini_file_uploader.get_document_manifest()
                
                cache_name = await self.gemini_cache_manager.create_cache_from_files_async(
                    uploaded_files=uploaded_files,
                    document_manifest=document_manifest,
                    ttl_seconds=7200,  # 2 hours
                    model=GeminiModels.DEFAULT_LLM,  # Must match LLM model
                )
                if cache_name:
                    getattr(self.llm, 'set_cached_content')(cache_name)
                    logger.info(f"LLM using file-based cache: {cache_name}")
                    logger.info(f"Document manifest:\n{document_manifest[:500]}...")
                    # Phase 8: Update manifest with cache timestamp
                    self.rag_manifest.update_cache_timestamp()
                else:
                    logger.warning("Cache creation returned empty cache_name")
            except Exception as e:
                logger.error(f"Failed to create file-based cache: {e}")
                # Fallback to chunk-based cache
                logger.info("Falling back to chunk-based cache...")
                try:
                    cache_name = await self.gemini_cache_manager.create_cache_from_context_async(
                        self.context_manager,
                        ttl_seconds=7200,
                        model=GeminiModels.DEFAULT_LLM,
                    )
                    if cache_name:
                        getattr(self.llm, 'set_cached_content')(cache_name)
                        logger.info(f"LLM using fallback chunk cache: {cache_name}")
                except Exception as fallback_err:
                    logger.error(f"Fallback cache also failed: {fallback_err}")
        else:
            logger.warning("No files uploaded to Gemini - cannot create file-based cache. Trying chunk-based cache...")
            try:
                cache_name = await self.gemini_cache_manager.create_cache_from_context_async(
                    self.context_manager,
                    ttl_seconds=7200,
                    model=GeminiModels.DEFAULT_LLM,
                )
                if cache_name:
                    getattr(self.llm, 'set_cached_content')(cache_name)
                    logger.info(f"LLM using chunk-based cache: {cache_name}")
            except Exception as e:
                logger.error(f"Chunk-based cache creation failed: {e}")

    async def _handle_calibrate_voice(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle CALIBRATE_VOICE message."""
        if not self.speaker_recognizer:
            # Try to get from warmer
            if self.model_warmer.wait_for_ready(timeout=2.0):
                models = self.model_warmer.get_models()
                if models.speaker_recognizer:
                    self.speaker_recognizer = cast(SpeakerRecognizer, models.speaker_recognizer)
            
            # If still None, try synchronous load
            if not self.speaker_recognizer:
                try:
                    logger.info("Loading SpeakerRecognizer synchronously for calibration")
                    self.speaker_recognizer = SpeakerRecognizer()
                except Exception as e:
                    logger.error(f"Failed to load SpeakerRecognizer: {e}")

        if not self.speaker_recognizer:
            error_msg = create_error_message(
                "Speaker recognizer not initialized",
                code="ERR_COMPONENT_NOT_READY"
            )
            await websocket.send(error_msg.to_json())
            return

        self.session_state.status = SessionStatus.CALIBRATING
        
        status_msg = create_status_message(SessionStatus.CALIBRATING)
        await websocket.send(status_msg.to_json())

        try:
            data = message.data or {}
            audio_b64 = data.get("audioData")
            
            if not audio_b64:
                raise ValueError("No audioData provided")
                
            audio_bytes = base64.b64decode(audio_b64)
            
            audio_chunk = np.frombuffer(audio_bytes, dtype=np.int16)
            
            loop = asyncio.get_running_loop()
            embedding = await loop.run_in_executor(
                None, 
                self.speaker_recognizer.create_embedding, 
                audio_chunk
            )
            
            self.session_state.user_embedding = embedding
            self.session_state.voice_calibrated = True
            
            logger.info("Voice calibration completed successfully")
            
            self.session_state.status = SessionStatus.IDLE
            status_msg = create_status_message(SessionStatus.IDLE)
            await websocket.send(status_msg.to_json())

        except Exception as e:
            logger.error(f"Voice calibration failed: {e}")
            self.session_state.status = SessionStatus.IDLE
            
            error_msg = create_error_message(
                f"Calibration failed: {str(e)}",
                code="ERR_CALIBRATION_FAILED"
            )
            await websocket.send(error_msg.to_json())
            
            status_msg = create_status_message(SessionStatus.IDLE)
            await websocket.send(status_msg.to_json())

    async def _handle_manual_question(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle MANUAL_QUESTION message."""
        data = message.data or {}
        question = data.get("question", "")

        if not question:
            error_msg = create_error_message(
                "Question is required",
                code="ERR_NO_QUESTION"
            )
            await websocket.send(error_msg.to_json())
            return

        logger.info(f"Manual question received: {question[:50]}...")

        self.session_state.status = SessionStatus.PROCESSING
        status_msg = create_status_message(SessionStatus.PROCESSING)
        await websocket.send(status_msg.to_json())

        # Phase 4E: Trigger coaching pipeline (story recall & structure suggestion)
        q_type = classify_question(question)
        if q_type in ("behavioral", "interview_question", "intro", "weakness", "conflict", "leadership"):
            asyncio.create_task(self._recall_and_suggest_story(question, q_type))
        asyncio.create_task(self._suggest_structure(question, q_type))

        retrieval_results = []
        context_chunks = []
        rag_confidence = ConfidenceLevel.LOW

        # Phase 5: Cache-First Architecture
        # If Gemini Cache is active, we SKIP RAG entirely because the LLM
        # already has the full document context in its memory window.
        has_cache = False
        if self.llm and hasattr(self.llm, 'has_cached_content') and self.llm.has_cached_content():
            has_cache = True
            logger.info("Using Gemini Context Cache - Skipping RAG retrieval")
        
        # Only perform RAG if no cache is active
        if self.rag_engine and not has_cache:
            try:
                # Use async retrieval to prevent blocking the WebSocket loop (Phase 5 fix)
                # Fallback to sync method if async not available (backward compat)
                if hasattr(self.rag_engine, 'retrieve_async'):
                    retrieval_results = await self.rag_engine.retrieve_async(question, limit=5)
                else:
                    logger.warning("RAG engine missing retrieve_async, using blocking call")
                    retrieval_results = self.rag_engine.retrieve(question, limit=5)
                    
                logger.info(f"Retrieved {len(retrieval_results)} chunks for question")
                context_chunks = [r.text for r in retrieval_results]

                if retrieval_results:
                    rag_confidence = self._confidence_from_string(retrieval_results[0].confidence)

                for i, r in enumerate(retrieval_results):
                    logger.debug(f"Chunk {i}: {r.confidence} ({r.distance:.2f}) - {r.text[:50]}...")
            except Exception as e:
                logger.error(f"RAG retrieval failed: {e}")

        if self.llm:
            try:
                # Signal start of answer
                await self.broadcast(Message(type=MessageType.ANSWER_START))
                
                context_str = "\n\n".join(context_chunks)
                # Collect full answer for history
                full_answer_parts: List[str] = []
                
                # Check for empty context
                if not context_chunks and "hiring manager" in self.session_state.conversation_history:
                     # This check is illustrative; actual fix involves clearing persistent state
                     pass 

                # Explicitly add Source Metadata to context string if RAG is used
                # This helps LLM distinguish between "Resume" vs "Interviewer Info"
                if context_chunks and retrieval_results:
                    context_parts = []
                    for res in retrieval_results:
                        meta = res.metadata or {}
                        doc_type = meta.get("document_type", "unknown")
                        source = meta.get("source", "unknown")
                        context_parts.append(f"Source: {source} ({doc_type})\n{res.text}")
                    context_str = "\n\n".join(context_parts)

                async for chunk in self.llm.generate_response(
                    question, context_str, self.session_state.conversation_history
                ):
                    full_answer_parts.append(chunk)
                    answer_msg = create_answer_chunk_message(
                        chunk=chunk,
                        complete=False
                    )
                    await websocket.send(answer_msg.to_json())

                await websocket.send(create_answer_chunk_message(
                    chunk="",
                    complete=True,
                    confidence=rag_confidence
                ).to_json())
                
                # Append Q&A to conversation history for future context
                full_answer = "".join(full_answer_parts)
                self.session_state.conversation_history.append({
                    "role": "user",
                    "content": question
                })
                self.session_state.conversation_history.append({
                    "role": "assistant",
                    "content": full_answer
                })
                logger.debug(f"Added manual Q&A to history. Total exchanges: {len(self.session_state.conversation_history) // 2}")
                
            except Exception as e:
                logger.error(f"LLM generation failed: {e}")
                error_msg = create_error_message(
                    f"Failed to generate answer: {e}",
                    code="ERR_LLM_GENERATION"
                )
                await websocket.send(error_msg.to_json())
        else:
             error_msg = create_error_message(
                "LLM not initialized",
                code="ERR_LLM_NOT_READY"
            )
             await websocket.send(error_msg.to_json())

        self.session_state.status = SessionStatus.LISTENING
        status_msg = create_status_message(SessionStatus.LISTENING)
        await websocket.send(status_msg.to_json())

    # Session History Handlers (Phase 3: STORY-039)

    async def _handle_list_sessions(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle LIST_SESSIONS message - returns paginated session list."""
        data = message.data or {}
        limit = data.get("limit", 20)
        offset = data.get("offset", 0)
        
        try:
            sessions = self.session_store.list_sessions(limit=limit, offset=offset)
            session_summaries = [self._session_summary_to_dict(s) for s in sessions]
            
            response = create_session_list_message(
                sessions=session_summaries,
                total=len(session_summaries),
                has_more=len(sessions) == limit
            )
            await websocket.send(response.to_json())
            logger.info(f"Listed {len(sessions)} sessions (offset={offset}, limit={limit})")
        except Exception as e:
            logger.error(f"Failed to list sessions: {e}")
            error_msg = create_error_message(
                f"Failed to list sessions: {e}",
                code="ERR_LIST_SESSIONS"
            )
            await websocket.send(error_msg.to_json())

    async def _handle_load_session(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle LOAD_SESSION message - returns full session data."""
        data = message.data or {}
        session_id = data.get("sessionId")
        
        if not session_id:
            error_msg = create_error_message(
                "sessionId is required",
                code="ERR_MISSING_SESSION_ID"
            )
            await websocket.send(error_msg.to_json())
            return
        
        try:
            session = self.session_store.get_session(session_id)
            if not session:
                error_msg = create_error_message(
                    f"Session not found: {session_id}",
                    code="ERR_SESSION_NOT_FOUND"
                )
                await websocket.send(error_msg.to_json())
                return
            
            response = create_session_data_message(self._session_to_full_dict(session))
            await websocket.send(response.to_json())
            logger.info(f"Loaded session: {session_id}")
        except Exception as e:
            logger.error(f"Failed to load session: {e}")
            error_msg = create_error_message(
                f"Failed to load session: {e}",
                code="ERR_LOAD_SESSION"
            )
            await websocket.send(error_msg.to_json())

    async def _handle_export_session(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle EXPORT_SESSION message - returns formatted session content."""
        data = message.data or {}
        session_id = data.get("sessionId")
        export_format = data.get("format", "md")
        
        if not session_id:
            error_msg = create_error_message(
                "sessionId is required",
                code="ERR_MISSING_SESSION_ID"
            )
            await websocket.send(error_msg.to_json())
            return
        
        try:
            session = self.session_store.get_session(session_id)
            if not session:
                error_msg = create_error_message(
                    f"Session not found: {session_id}",
                    code="ERR_SESSION_NOT_FOUND"
                )
                await websocket.send(error_msg.to_json())
                return
            
            if export_format == "json":
                content = self._export_session_as_json(session)
            elif export_format == "txt":
                content = SessionExporter.export(session, ExportFormat.TEXT)
            else:
                content = self._export_session_as_markdown(session)
                export_format = "md"
            
            response = create_session_export_message(content=content, format=export_format)
            await websocket.send(response.to_json())
            logger.info(f"Exported session {session_id} as {export_format}")
        except Exception as e:
            logger.error(f"Failed to export session: {e}")
            error_msg = create_error_message(
                f"Failed to export session: {e}",
                code="ERR_EXPORT_SESSION"
            )
            await websocket.send(error_msg.to_json())

    async def _handle_delete_session(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle DELETE_SESSION message - deletes session and confirms."""
        data = message.data or {}
        session_id = data.get("sessionId")
        
        if not session_id:
            error_msg = create_error_message(
                "sessionId is required",
                code="ERR_MISSING_SESSION_ID"
            )
            await websocket.send(error_msg.to_json())
            return
        
        try:
            success = self.session_store.delete_session(session_id)
            response = create_session_deleted_message(session_id=session_id, success=success)
            await websocket.send(response.to_json())
            
            if success:
                logger.info(f"Deleted session: {session_id}")
            else:
                logger.warning(f"Session not found for deletion: {session_id}")
        except Exception as e:
            logger.error(f"Failed to delete session: {e}")
            error_msg = create_error_message(
                f"Failed to delete session: {e}",
                code="ERR_DELETE_SESSION"
            )
            await websocket.send(error_msg.to_json())

    async def _handle_prepare_interview(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """
        Handle PREPARE_INTERVIEW message - generates pre-interview preparation summary.
        
        Phase 3B: STORY-047
        """
        logger.info("Starting interview preparation...")
        
        data = message.data or {}
        if not self.llm and data.get("apiKeys"):
            try:
                config = ProviderConfig.from_dict(data)
                factory = ProviderFactory(config)
                self.llm = factory.get_llm_provider()
                logger.info("Initialized LLM provider for preparation")
            except Exception as e:
                logger.warning(f"Could not initialize LLM from preparation data: {e}")
        
        chunks = self.context_manager.get_all_chunks()
        
        if not chunks:
            summary = """## Interview Preparation

**No documents uploaded.** Upload your resume and job description for personalized preparation.

To get started:
1. Upload your resume/CV
2. Upload the job description
3. Optionally add company information
4. Click "Prepare for Interview" again"""
            
            response = create_preparation_ready_message(summary)
            await websocket.send(response.to_json())
            return
        
        # Build preparation prompt from context
        try:
            preparation_prompt = self._build_preparation_prompt(chunks)
            
            if not self.llm:
                # Fallback if no LLM available
                summary = self._generate_fallback_preparation(chunks)
            else:
                # Generate preparation using LLM
                summary = await self._generate_preparation_with_llm(preparation_prompt)
            
            # Store in session state
            self.session_state.preparation_summary = summary
            
            response = create_preparation_ready_message(summary)
            await websocket.send(response.to_json())
            logger.info("Interview preparation complete")
            
        except Exception as e:
            logger.error(f"Failed to generate preparation: {e}")
            error_msg = create_error_message(
                f"Failed to generate preparation: {e}",
                code="ERR_PREPARATION_FAILED"
            )
            await websocket.send(error_msg.to_json())
    
    def _build_preparation_prompt(self, chunks: list) -> str:
        """Build prompt for preparation generation from context chunks."""
        # Group chunks by source/type
        resume_text = []
        jd_text = []
        company_text = []
        other_text = []
        
        for chunk in chunks:
            meta = chunk.metadata or {}
            source = meta.get("source", "").lower()
            doc_type = meta.get("document_type", "").lower()
            
            if "resume" in source or "cv" in source or doc_type == "resume":
                resume_text.append(chunk.text)
            elif "job" in source or "jd" in source or doc_type == "job_description":
                jd_text.append(chunk.text)
            elif "company" in source or "about" in source or doc_type == "company_info":
                company_text.append(chunk.text)
            else:
                other_text.append(chunk.text)
        
        sections = []
        
        if resume_text:
            sections.append(f"## CANDIDATE BACKGROUND\n{' '.join(resume_text[:5])}")
        
        if jd_text:
            sections.append(f"## ROLE REQUIREMENTS\n{' '.join(jd_text[:5])}")
        
        if company_text:
            sections.append(f"## COMPANY CONTEXT\n{' '.join(company_text[:3])}")
        
        if other_text and not (resume_text or jd_text):
            sections.append(f"## ADDITIONAL CONTEXT\n{' '.join(other_text[:3])}")
        
        return f"""Based on the following documents, prepare a comprehensive interview briefing:

{chr(10).join(sections)}

Generate a structured preparation summary with:
1. **Key Talking Points**: 3-5 points that align the candidate's experience with role requirements
2. **Potential Challenges**: 2-3 areas where the candidate may need to address gaps
3. **Company-Specific Insights**: 2-3 talking points that reference company values/products (if available)
4. **STAR Story Suggestions**: 2-3 specific experiences from the resume that could be used for behavioral questions
5. **Questions to Ask**: 2-3 intelligent questions the candidate could ask

Keep the briefing concise and actionable. Use bullet points and markdown formatting."""
    
    async def _generate_preparation_with_llm(self, prompt: str) -> str:
        """Generate preparation summary using LLM."""
        if not self.llm:
            raise RuntimeError("LLM not available")
        
        full_response_parts: List[str] = []
        
        async for chunk in self.llm.generate_response(
            prompt,
            "",  # No additional context
            []   # No conversation history
        ):
            full_response_parts.append(chunk)
        
        return "".join(full_response_parts)
    
    def _generate_fallback_preparation(self, chunks: list) -> str:
        """Generate basic preparation without LLM."""
        doc_count = len(chunks)
        
        return f"""## Interview Preparation Summary

**Documents Analyzed**: {doc_count} context chunks loaded

### Key Points from Your Documents

Based on the uploaded documents, here are the main topics to prepare:

{chr(10).join([f"- {chunk.text[:100]}..." for chunk in chunks[:5]])}

### General Preparation Tips

1. **Review your experience** aligned with the role requirements
2. **Prepare STAR stories** for behavioral questions
3. **Research the company** culture and recent news
4. **Prepare thoughtful questions** to ask the interviewer

*Note: For personalized insights, ensure an LLM provider is configured.*"""

    # Phase 5: Answer Enhancement Handler
    
    async def _handle_enhance_answer(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """
        Handle ENHANCE_ANSWER message - improves an existing answer.
        
        Phase 5: Answer Enhancement Feature
        
        Expected data:
        {
            "enhancementType": "add_detail" | "make_specific" | "suggest_star" | "adjust_tone" | "shorten",
            "originalQuestion": "...",
            "originalAnswer": "...",
            "tonePreference": "confident" | "humble" (optional, for adjust_tone)
        }
        """
        data = message.data or {}
        enhancement_type_str = data.get("enhancementType", "add_detail")
        original_question = data.get("originalQuestion", "")
        original_answer = data.get("originalAnswer", "")
        tone_preference = data.get("tonePreference", "confident")
        
        if not original_question or not original_answer:
            error_msg = create_error_message(
                "originalQuestion and originalAnswer are required",
                code="ERR_MISSING_DATA"
            )
            await websocket.send(error_msg.to_json())
            return
        
        try:
            enhancement_type = EnhancementType(enhancement_type_str)
        except ValueError:
            error_msg = create_error_message(
                f"Invalid enhancement type: {enhancement_type_str}",
                code="ERR_INVALID_ENHANCEMENT_TYPE"
            )
            await websocket.send(error_msg.to_json())
            return
        
        if not self.llm:
            error_msg = create_error_message(
                "LLM not initialized",
                code="ERR_LLM_NOT_READY"
            )
            await websocket.send(error_msg.to_json())
            return
        
        logger.info(f"Enhancing answer with type: {enhancement_type.value}")
        
        # Signal start of enhanced answer
        start_msg = create_enhanced_answer_start_message(enhancement_type, original_question)
        await websocket.send(start_msg.to_json())
        
        # Cancel any existing enhancement task
        if self._enhancement_task and not self._enhancement_task.done():
            self._enhancement_task.cancel()
            try:
                await self._enhancement_task
            except asyncio.CancelledError:
                pass
        
        # Run enhancement as a cancellable task
        self._enhancement_task = asyncio.create_task(
            self._run_enhancement(
                websocket, enhancement_type, original_question, 
                original_answer, tone_preference
            )
        )
        
        try:
            await self._enhancement_task
        except asyncio.CancelledError:
            logger.info("Enhancement was cancelled")
            # Send completion message to signal cancellation
            complete_msg = create_enhanced_answer_complete_message(enhancement_type, success=False)
            await websocket.send(complete_msg.to_json())
        finally:
            self._enhancement_task = None
    
    async def _run_enhancement(
        self,
        websocket: ServerConnection,
        enhancement_type: EnhancementType,
        original_question: str,
        original_answer: str,
        tone_preference: str
    ) -> None:
        """Run the actual enhancement streaming. Can be cancelled."""
        # Type assertion (already checked in caller)
        assert self.llm is not None, "LLM must be initialized before calling _run_enhancement"
        
        try:
            # Build enhancement prompt based on type
            enhancement_prompt = self._build_enhancement_prompt(
                enhancement_type=enhancement_type,
                question=original_question,
                answer=original_answer,
                tone_preference=tone_preference
            )
            
            # Get additional context if needed (for add_detail)
            context_str = ""
            if enhancement_type == EnhancementType.ADD_DETAIL and self.rag_engine:
                try:
                    results = self.rag_engine.retrieve(original_question, limit=8)  # Higher limit
                    context_str = "\n\n".join([r.text for r in results])
                except Exception as e:
                    logger.warning(f"RAG retrieval for enhancement failed: {e}")
            
            # Stream enhanced answer
            async for chunk in self.llm.generate_response(
                enhancement_prompt, context_str, []
            ):
                chunk_msg = create_enhanced_answer_chunk_message(chunk=chunk, complete=False)
                await websocket.send(chunk_msg.to_json())
            
            # Signal completion
            complete_msg = create_enhanced_answer_complete_message(enhancement_type, success=True)
            await websocket.send(complete_msg.to_json())
            
            logger.info(f"Enhancement complete: {enhancement_type.value}")
            
        except asyncio.CancelledError:
            # Re-raise to let the caller handle it
            raise
        except Exception as e:
            logger.error(f"Enhancement failed: {e}")
            error_msg = create_error_message(
                f"Enhancement failed: {e}",
                code="ERR_ENHANCEMENT_FAILED"
            )
            await websocket.send(error_msg.to_json())
    
    async def _handle_cancel_enhancement(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle CANCEL_ENHANCEMENT message - stops ongoing enhancement."""
        if self._enhancement_task and not self._enhancement_task.done():
            self._enhancement_task.cancel()
            try:
                await self._enhancement_task
            except asyncio.CancelledError:
                pass
            self._enhancement_task = None
            logger.info("Enhancement cancelled by user request")
        else:
            logger.debug("No enhancement task to cancel")
    
    def _build_enhancement_prompt(
        self,
        enhancement_type: EnhancementType,
        question: str,
        answer: str,
        tone_preference: str = "confident"
    ) -> str:
        """Build the prompt for answer enhancement based on type."""
        
        base_context = f"""Original Question: {question}

Current Answer: {answer}

"""
        
        if enhancement_type == EnhancementType.ADD_DETAIL:
            return base_context + """Enhance this answer by:
1. Adding more specific details from the provided context
2. Including relevant examples or metrics if available
3. Expanding on key points that could use more depth
4. Maintaining the same overall structure and tone

Provide an improved, more detailed version of the answer. Keep it natural and conversational."""

        elif enhancement_type == EnhancementType.MAKE_SPECIFIC:
            return base_context + """Make this answer more specific by:
1. Adding concrete numbers, percentages, or metrics where possible
2. Including specific examples with names/dates/technologies
3. Quantifying achievements and impact
4. Replacing vague statements with specific details

Provide a version with more specifics. If you don't have exact data, suggest realistic placeholder specifics the candidate should fill in [like this]."""

        elif enhancement_type == EnhancementType.SUGGEST_STAR:
            return base_context + """Transform this answer into a STAR story format:
1. **Situation**: Set the context with specific details
2. **Task**: What was your specific responsibility?
3. **Action**: What steps did YOU specifically take?
4. **Result**: What was the measurable outcome?

Rewrite the answer following the STAR structure. Make it compelling and specific."""

        elif enhancement_type == EnhancementType.ADJUST_TONE:
            tone_desc = "more confident and assertive, using strong action verbs" if tone_preference == "confident" else "more humble and collaborative, acknowledging team contributions"
            return base_context + f"""Adjust the tone of this answer to be {tone_desc}.

Rewrite the answer maintaining the same information but with the adjusted tone. Keep it natural and authentic."""

        elif enhancement_type == EnhancementType.SHORTEN:
            return base_context + """Shorten this answer while keeping the key points:
1. Remove redundant phrases and filler words
2. Condense to the most impactful statements
3. Keep the strongest points, remove weaker ones
4. Aim for ~50% of the original length

Provide a concise, punchy version that hits the key points quickly."""

        else:
            return base_context + "Improve this answer while maintaining its core message."

    # Phase 5: Document Type Inference Handler
    
    async def _handle_infer_document_types(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """Handle INFER_DOCUMENT_TYPES message."""
        data = message.data or {}
        files = data.get("files", [])
        
        # Lazy Initialization: If API keys are present in this message but providers aren't ready,
        # initialize them now. This handles the case where users add files before "Starting Session".
        has_keys = "apiKeys" in data or "apiKey" in data
        if has_keys and not self.llm:
            try:
                logger.info("Auto-initializing providers for document inference...")
                # Ensure apiKeys structure matches ProviderConfig expectation
                if "apiKeys" not in data and "apiKey" in data:
                     data["apiKeys"] = {"gemini": data["apiKey"]}
                     
                config = ProviderConfig.from_dict(data)
                self.provider_factory = ProviderFactory(config)
                
                # We need LLM for classification
                self.llm = self.provider_factory.get_llm_provider()
                
                # Set the LLM on the classifier immediately
                if self.llm:
                    self.document_classifier.set_llm_provider(self.llm)
                
                # Also init other providers if possible, to be ready for session start
                if not self.stt:
                    self.stt = self.provider_factory.get_stt_provider()
                    
                logger.info("Providers initialized for inference")
            except Exception as e:
                logger.warning(f"Failed to auto-init providers for inference: {e}")

        if not files:
            error_msg = create_error_message(
                "No files provided for type inference",
                code="ERR_NO_FILES"
            )
            await websocket.send(error_msg.to_json())
            return
        
        logger.info(f"Inferring document types for {len(files)} file(s)")
        
        # Ensure classifier has LLM if available (use Flash for speed)
        if self.llm and not self.document_classifier._llm:
            # Try to get a fast LLM for classification
            if self.provider_factory:
                try:
                    # Use the existing LLM - it will still be fast enough
                    self.document_classifier.set_llm_provider(self.llm)
                    logger.info("Set LLM provider for document classifier")
                except Exception as e:
                    logger.warning(f"Failed to set classifier LLM: {e}")
        
        suggestions = []
        
        for file_data in files:
            file_id = file_data.get("id", file_data.get("filename", "unknown"))
            filename = file_data.get("filename", "unknown")
            content_b64 = file_data.get("content", "")
            
            try:
                # Extract text content from file
                text_content = await self._extract_text_for_classification(
                    content_b64=content_b64,
                    filename=filename
                )
                
                # Classify using LLM or fallback to heuristics
                result = await self.document_classifier.classify(
                    filename=filename,
                    text_content=text_content,
                    fallback_type="custom"
                )
                
                suggestions.append({
                    "id": file_id,
                    "filename": filename,
                    **result.to_dict()
                })
                
                logger.info(f"Classified '{filename}' as {result.document_type} ({result.confidence:.0%})")
                
            except Exception as e:
                logger.error(f"Failed to classify '{filename}': {e}")
                # Add fallback suggestion
                suggestions.append({
                    "id": file_id,
                    "filename": filename,
                    "documentType": "custom",
                    "confidence": 0.3,
                    "reason": f"Classification failed: {str(e)[:50]}"
                })
        
        # Send suggestions back to frontend
        response_msg = create_document_type_suggestions_message(suggestions)
        await websocket.send(response_msg.to_json())
        
        logger.info(f"Sent {len(suggestions)} document type suggestion(s)")
    
    async def _extract_text_for_classification(
        self,
        content_b64: str,
        filename: str
    ) -> str:
        """
        Extract text content from a base64-encoded file for classification.
        
        Supports: .txt, .md, .docx, .pdf (first page via fallback)
        Returns first ~2000 chars for classification.
        """
        import io
        import docx
        
        try:
            content = base64.b64decode(content_b64)
            ext = os.path.splitext(filename)[1].lower()
            
            # Plain text files
            if ext in ['.txt', '.md', '.html', '.json']:
                text = content.decode('utf-8', errors='ignore')
                return text[:2000]
            
            # DOCX files
            if ext == '.docx':
                doc = docx.Document(io.BytesIO(content))
                paragraphs = [para.text for para in doc.paragraphs]
                text = '\n'.join(paragraphs)
                return text[:2000]
            
            # PDF files - basic text extraction
            if ext == '.pdf':
                try:
                    import pypdf
                    reader = pypdf.PdfReader(io.BytesIO(content))
                    text_parts = []
                    for page in reader.pages[:3]:  # First 3 pages max
                        text_parts.append(page.extract_text() or "")
                        if len(''.join(text_parts)) > 2000:
                            break
                    return ''.join(text_parts)[:2000]
                except ImportError:
                    logger.warning("pypdf not available, using filename heuristics for PDF")
                    return ""  # Will trigger filename-based fallback
                except Exception as e:
                    logger.warning(f"PDF extraction failed: {e}")
                    return ""
            
            # Unknown format - try as text
            try:
                text = content.decode('utf-8', errors='ignore')
                return text[:2000]
            except Exception:
                return ""
                
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            return ""

    # Session History Helper Methods

    def _session_summary_to_dict(self, session) -> dict:
        """Convert a SessionSummary or SessionData object to a summary dict."""
        from src.storage.session_store import SessionSummary, SessionData
        
        # Determine counts based on object type
        if isinstance(session, SessionSummary):
            trans_count = session.transcription_count
            ans_count = session.answer_count
        else:
            # Assume SessionData
            trans_count = len(getattr(session, "transcriptions", []))
            ans_count = len(getattr(session, "answers", []))
            
        return {
            "id": session.id,
            "startedAt": int(session.started_at.timestamp() * 1000) if session.started_at else None,
            "endedAt": int(session.ended_at.timestamp() * 1000) if session.ended_at else None,
            "contextFiles": session.context_files,
            "transcriptionCount": trans_count,
            "answerCount": ans_count
        }

    def _session_to_full_dict(self, session) -> dict:
        """Convert a Session object to a full dict with all data."""
        from src.storage.session_store import SessionData
        s: SessionData = session
        return {
            "id": s.id,
            "startedAt": int(s.started_at.timestamp() * 1000) if s.started_at else None,
            "endedAt": int(s.ended_at.timestamp() * 1000) if s.ended_at else None,
            "contextFiles": s.context_files,
            "transcriptions": s.transcriptions,
            "answers": s.answers
        }

    def _export_session_as_markdown(self, session) -> str:
        """Export a session as formatted markdown using SessionExporter."""
        return SessionExporter.export(session, ExportFormat.MARKDOWN)

    def _export_session_as_json(self, session) -> str:
        """Export a session as JSON string using SessionExporter."""
        return SessionExporter.export(session, ExportFormat.JSON)

    async def broadcast(self, message: Message) -> None:
        """
        Broadcast a message to all connected clients.

        Args:
            message: The message to broadcast
        """
        if not self.clients:
            return

        await asyncio.gather(
            *[client.send(message.to_json()) for client in self.clients],
            return_exceptions=True
        )

    async def _start_audio_processing(self) -> None:
        """Initialize and start audio processing components."""
        # STT is already initialized in _handle_start_session via factory
        if not self.stt:
             logger.error("STT provider not initialized")
             raise RuntimeError("STT provider not initialized")
        
        if not self.model_warmer.wait_for_ready(timeout=2.0):
             logger.warning("Models not ready after timeout")
        
        models = self.model_warmer.get_models()
        if models.is_ready and models.vad_processor and models.speaker_recognizer:
             self.vad = cast(VADProcessor, models.vad_processor)
             self.vad.reset()
             self.speaker_recognizer = cast(SpeakerRecognizer, models.speaker_recognizer)
        else:
             logger.info("Initializing models synchronously")
             self.vad = VADProcessor()
             self.speaker_recognizer = SpeakerRecognizer()

        self.noise_reducer = NoiseReducer(enabled=True)
        self.audio_capture = AudioCapture()
        
        await self.audio_capture.start_capture()
        self._audio_task = asyncio.create_task(self._audio_loop())
        logger.info("Audio processing started")

    async def _init_rag_background(self, api_key: str) -> None:
        """
        Initialize RAG components in background.
        
        This runs asynchronously to avoid blocking audio processing.
        Session continues to work even if RAG initialization fails.
        """
        try:
            logger.info("Initializing RAG in background...")
            
            # Run blocking VectorStore initialization in executor
            loop = asyncio.get_event_loop()
            vector_store = await loop.run_in_executor(
                None, 
                lambda: VectorStore(api_key=api_key)
            )
            self.vector_store = vector_store
            self.rag_engine = EnhancedRAGEngine(vector_store, context_manager=self.context_manager)
            
            # Phase 4D: Initialize Speculative Retriever
            self.speculative_retriever = SpeculativeRetriever(self.rag_engine)
            
            # Phase 4E: Initialize Story Recaller
            if self.memory_store:
                self.story_recaller = StoryRecaller(self.memory_store, self.vector_store)
                # Warm up stories in background
                self._create_background_task(self.story_recaller.warm_up())
            
            self.structure_suggester = StructureSuggester()
            
            # Add pre-loaded context chunks
            pre_loaded_chunks = self.context_manager.get_all_chunks()
            if pre_loaded_chunks:
                logger.info(f"Adding {len(pre_loaded_chunks)} pre-loaded chunks to vector store (background)")
                chunk_texts = [c.text for c in pre_loaded_chunks]
                chunk_metas = [c.metadata for c in pre_loaded_chunks]
                try:
                    # Run blocking add_documents in executor
                    def add_docs():
                        vector_store.add_documents(chunk_texts, metadatas=chunk_metas)
                    await loop.run_in_executor(None, add_docs)
                    logger.info("RAG initialization complete - context loaded")
                except Exception as e:
                    logger.warning(f"Failed to add pre-loaded chunks (RAG degraded): {e}")
            else:
                logger.info("RAG initialization complete - no pre-loaded chunks")
                    
        except Exception as e:
            logger.warning(f"RAG initialization failed (session continues without RAG): {e}")
            # Don't raise - session continues without RAG support
            self.vector_store = None
            self.rag_engine = None

    async def _stop_audio_processing(self) -> None:
        """Stop audio processing components."""
        if self._audio_task:
            self._audio_task.cancel()
            try:
                await self._audio_task
            except asyncio.CancelledError:
                pass
            self._audio_task = None
            
        if self.audio_capture:
            await self.audio_capture.stop_capture()
            self.audio_capture = None
            
        self.stt = None
        self.vad = None
        self.noise_reducer = None
        logger.info("Audio processing stopped")

    async def _audio_loop(self) -> None:
        """
        Main audio processing loop.
        
        Pipeline: Audio → VAD → NoiseReducer → STT → Diarization → RAG → LLM → UI
        
        - User speech: transcribed and displayed (filtered from RAG+LLM)
        - Interviewer speech: transcribed, then triggers RAG retrieval and LLM generation
        """
        import time
        if not self.audio_capture or not self.vad or not self.stt:
            logger.error("Audio components not initialized")
            return

        logger.info("Starting audio processing loop")
        
        last_speculation_time = 0.0
        
        try:
            async for chunk in self.audio_capture.get_audio_stream():
                # Phase 7: Route audio to streaming STT for real-time transcription
                # This runs in parallel with VAD-based batch STT
                if self.streaming_stt_manager and self.streaming_stt_manager.is_active:
                    # Send audio chunk to streaming provider (non-blocking)
                    await self.streaming_stt_manager.send_audio(chunk)
                
                segments = await self.vad.process_chunk(chunk)
                
                for segment in segments:
                    # Reset speculation state on segment completion
                    if self.speculative_retriever:
                        self.speculative_retriever.reset()
                    await self._process_speech_segment(segment)
                
                # Phase 4D: Speculative Retrieval Trigger
                # Check periodically if we should speculatively retrieve for ongoing speech
                now = time.time()
                if (self.speculative_retriever and 
                    self.vad.is_speaking and 
                    self.vad.current_duration > 2.0 and 
                    now - last_speculation_time > 2.0):
                    
                    # Run in background to avoid blocking audio loop
                    self._create_background_task(self._run_speculative_cycle())
                    last_speculation_time = now
                        
        except asyncio.CancelledError:
            logger.info("Audio loop cancelled")
            raise
        except Exception as e:
            logger.error(f"Fatal error in audio loop: {e}")
            error_msg = create_error_message(
                f"Audio processing error: {e}",
                code="ERR_AUDIO_LOOP"
            )
            await self.broadcast(error_msg)
            self.session_state.status = SessionStatus.IDLE
            await self.broadcast(create_status_message(SessionStatus.IDLE))

    # =========================================================================
    # Phase 7: Streaming STT Callbacks
    # These handle real-time transcription and semantic endpointing events
    # =========================================================================
    
    def _on_streaming_interim(self, text: str, speaker: Speaker) -> None:
        """Handle interim transcription from streaming STT."""
        if not text:
            return
        
        # Broadcast interim transcription to UI (non-blocking)
        asyncio.create_task(self._broadcast_interim_transcription(text, speaker))
    
    async def _broadcast_interim_transcription(self, text: str, speaker: Speaker) -> None:
        """Broadcast interim transcription to UI."""
        import time
        interim_msg = create_interim_transcription_message(
            text=text,
            timestamp=time.time(),
            speaker=speaker
        )
        await self.broadcast(interim_msg)
    
    def _on_streaming_final(self, text: str, speaker: Speaker, confidence: float) -> None:
        """Handle final transcription from streaming STT (before end-of-turn)."""
        if not text:
            return
        
        logger.debug(f"Streaming final ({speaker.value}, conf={confidence:.2f}): {text[:50]}...")
        # Final transcripts are handled by end-of-turn or batch fallback
    
    def _on_streaming_end_of_turn(
        self, 
        text: str, 
        speaker: Speaker, 
        confidence: float, 
        endpointing_type: EndpointingType
    ) -> None:
        """
        Handle end-of-turn event from streaming STT provider.
        
        This is the semantic endpointing signal that indicates the speaker
        has finished their utterance. Can be acoustic (pause-based) or
        semantic (meaning-based from providers like AssemblyAI or OpenAI).
        
        Args:
            text: Final transcribed text for the turn
            speaker: Speaker identifier
            confidence: Confidence score (0.0-1.0)
            endpointing_type: ACOUSTIC or SEMANTIC
        """
        if not text:
            return
        
        logger.info(f"Streaming end-of-turn ({speaker.value}, {endpointing_type.value}, "
                   f"conf={confidence:.2f}): {text[:60]}...")
        
        # Fire the end-of-turn processing in background
        asyncio.create_task(self._process_streaming_end_of_turn(
            text, speaker, confidence, endpointing_type
        ))
    
    async def _process_streaming_end_of_turn(
        self,
        text: str,
        speaker: Speaker,
        confidence: float,
        endpointing_type: EndpointingType
    ) -> None:
        """
        Process end-of-turn event from streaming STT.
        
        This integrates with the UtteranceAccumulator in hybrid mode:
        - If confidence is high enough, bypass timing-based detection
        - Trigger question processing pipeline directly
        """
        import time
        pipeline_start = time.time()
        
        # Broadcast final transcription to UI
        transcription_msg = create_transcription_message(
            speaker=speaker,
            text=text,
            timestamp=pipeline_start,
            confidence=confidence
        )
        await self.broadcast(transcription_msg)
        logger.info(f"Streaming transcribed ({speaker.value}): {text[:50]}...")
        
        # Persist transcription to session history
        if self.session_persistence_enabled and self.session_state.persistent_session_id:
            try:
                self.session_store.add_transcription(
                    session_id=self.session_state.persistent_session_id,
                    speaker=speaker.value,
                    text=text,
                    timestamp=pipeline_start,
                    confidence=confidence
                )
            except Exception as e:
                logger.warning(f"Failed to persist streaming transcription: {e}")
        
        # Only process interviewer speech for answer generation
        if speaker != Speaker.INTERVIEWER:
            return
        
        # Phase 7: Integrate with UtteranceAccumulator in hybrid mode
        # High-confidence semantic endpointing can bypass timing-based detection
        if (self.accumulation_enabled and 
            hasattr(self, 'utterance_accumulator') and 
            self.utterance_accumulator):
            
            # Use the accumulator's streaming endpoint handler
            is_semantic = endpointing_type == EndpointingType.SEMANTIC
            complete_utterance = await self.utterance_accumulator.on_streaming_end_of_turn(
                text=text,
                speaker=speaker.value,
                confidence=confidence,
                is_semantic=is_semantic,
                timestamp=pipeline_start,
            )
            
            if complete_utterance is not None:
                # Streaming endpoint accepted - use the complete utterance
                text = complete_utterance.text
                logger.info(f"Streaming endpoint accepted ({complete_utterance.tier_used}): {text[:80]}...")
            else:
                # Low confidence - fall back to timing-based accumulation
                complete_utterance = await self.utterance_accumulator.add_segment(
                    text=text,
                    speaker=speaker.value,
                    timestamp=pipeline_start,
                    is_final=True
                )
                
                if complete_utterance is None:
                    # Still buffering
                    buffer = self.utterance_accumulator.get_buffer(speaker.value)
                    if buffer and self.utterance_accumulator.config.emit_accumulating_status:
                        preview = self.utterance_accumulator.get_buffer_preview(speaker.value)
                        accumulating_msg = create_accumulating_message(
                            speaker=speaker.value,
                            buffer_preview=preview,
                            segment_count=buffer.segment_count,
                            duration_s=buffer.duration_s
                        )
                        await self.broadcast(accumulating_msg)
                    return
                
                # Use accumulated text
                text = complete_utterance.text
                logger.info(f"Streaming accumulated complete: {text[:80]}...")
        
        # Question detection and answer generation
        if self.question_detection_enabled:
            is_question, q_confidence, q_type = await self.question_detector.is_actionable_question_async(
                text,
                self.session_state.conversation_history
            )
            logger.info(f"Streaming question detection: {q_type} (confidence={q_confidence:.2f})")
            
            if is_question and q_confidence >= self.question_confidence_threshold:
                await self._process_question_pipeline(text, q_type, pipeline_start)
            else:
                logger.info(f"Skipping answer for non-question ({q_type}): {text[:50]}...")
        else:
            await self._generate_answer_for_question(text, pipeline_start)
    
    def _on_streaming_error(self, error: Exception) -> None:
        """Handle streaming STT error."""
        logger.error(f"Streaming STT error: {error}")
        # Don't broadcast to UI - graceful degradation to batch STT

    async def _run_speculative_cycle(self) -> None:
        """
        Run a single cycle of speculative retrieval.
        Captures current audio buffer, transcribes it, and triggers retrieval.
        """
        import time
        if not self.stt or not self.speculative_retriever or not self.vad:
            return
            
        try:
            audio_buffer = self.vad.get_current_audio()
            # Need at least ~1s of audio (16000 samples * 2 bytes = 32000 bytes)
            if not audio_buffer or len(audio_buffer) < 32000:
                return
                
            # Quick transcription of partial buffer
            # Note: This relies on STT provider being able to handle short partials
            result = await self.stt.transcribe(audio_buffer)
            if result.text:
                # Phase 4D: Broadcast interim transcript
                # Calculate approx start time based on current duration
                start_time = time.time() - self.vad.current_duration
                interim_msg = create_interim_transcription_message(
                    text=result.text,
                    timestamp=start_time,
                    speaker=Speaker.INTERVIEWER
                )
                await self.broadcast(interim_msg)
                
                await self.speculative_retriever.on_interim_transcript(result.text)
                
        except Exception as e:
            logger.debug(f"Speculative cycle failed: {e}")

    async def _process_speech_segment(self, segment) -> None:
        """Process a speech segment: NoiseReducer → Speaker ID → STT → Broadcast → (RAG+LLM if Interviewer)."""
        import time
        pipeline_start = time.time()
        
        audio_for_stt = segment.audio
        if self.noise_reducer and self.noise_reducer.enabled:
            clean_audio = self.noise_reducer.reduce_noise(segment.audio)
            if isinstance(clean_audio, bytes):
                audio_for_stt = clean_audio
        
        speaker = self._identify_speaker(segment.audio)
        
        try:
            if not self.stt:
                logger.error("STT component is None during speech processing")
                return

            result = await self.stt.transcribe(audio_for_stt)
            text = result.text
            if not text:
                return
        except Exception as e:
            logger.error(f"STT error: {e}")
            return
        
        transcription_msg = create_transcription_message(
            speaker=speaker,
            text=text,
            timestamp=segment.start_time,
            confidence=segment.confidence
        )
        await self.broadcast(transcription_msg)
        logger.info(f"Transcribed ({speaker.value}): {text[:50]}...")
        
        # Phase 3: Persist transcription to session history
        if self.session_persistence_enabled and self.session_state.persistent_session_id:
            try:
                self.session_store.add_transcription(
                    session_id=self.session_state.persistent_session_id,
                    speaker=speaker.value,
                    text=text,
                    timestamp=segment.start_time,
                    confidence=segment.confidence
                )
            except Exception as e:
                logger.warning(f"Failed to persist transcription: {e}")
        
        if speaker == Speaker.INTERVIEWER:
            # Phase 6: Utterance Accumulation
            # Buffer segments from the same speaker until the utterance is semantically complete
            if self.accumulation_enabled and hasattr(self, 'utterance_accumulator'):
                import time as time_module
                complete_utterance = await self.utterance_accumulator.add_segment(
                    text=text,
                    speaker=speaker.value,
                    timestamp=time_module.time(),
                    is_final=True
                )
                
                if complete_utterance is None:
                    # Still buffering - send status update to UI
                    buffer = self.utterance_accumulator.get_buffer(speaker.value)
                    if buffer and self.utterance_accumulator.config.emit_accumulating_status:
                        preview = self.utterance_accumulator.get_buffer_preview(speaker.value)
                        accumulating_msg = create_accumulating_message(
                            speaker=speaker.value,
                            buffer_preview=preview,
                            segment_count=buffer.segment_count,
                            duration_s=buffer.duration_s
                        )
                        await self.broadcast(accumulating_msg)
                    logger.info(f"Accumulating: {text[:50]}... (buffering)")
                    return  # Don't process further until utterance is complete
                
                # Use the complete utterance text for question detection
                text = complete_utterance.text
                logger.info(f"Accumulated complete ({complete_utterance.completion_reason}): {text[:80]}...")
            
            # Phase 3: Question detection before answer generation
            if self.question_detection_enabled:
                is_question, confidence, q_type = await self.question_detector.is_actionable_question_async(
                    text,
                    self.session_state.conversation_history
                )
                logger.info(f"Question detection: {q_type} (confidence={confidence:.2f})")
                
                if is_question and confidence >= self.question_confidence_threshold:
                    # Phase 3C: Enhanced processing pipeline
                    await self._process_question_pipeline(text, q_type, pipeline_start)
                else:
                    logger.info(f"Skipping answer for non-question ({q_type}): {text[:50]}...")
            else:
                # Feature flag disabled - original behavior
                await self._generate_answer_for_question(text, pipeline_start)

    def _identify_speaker(self, audio: bytes) -> Speaker:
        """Identify speaker as User or Interviewer based on voice calibration."""
        if not self.session_state.voice_calibrated or self.session_state.user_embedding is None:
            return Speaker.INTERVIEWER
            
        if not self.speaker_recognizer:
            return Speaker.INTERVIEWER
            
        try:
            is_user = self.speaker_recognizer.verify_speaker(
                np.frombuffer(audio, dtype=np.int16),
                self.session_state.user_embedding
            )
            return Speaker.USER if is_user else Speaker.INTERVIEWER
        except Exception as e:
            logger.warning(f"Speaker verification failed: {e}, defaulting to Interviewer")
            return Speaker.INTERVIEWER

    async def _recall_and_suggest_story(self, question: str, q_type: str) -> None:
        """Find and broadcast relevant story match."""
        if not self.story_recaller:
            return
            
        try:
            match = await self.story_recaller.find_relevant_story(question, q_type)
            if match:
                msg = create_story_suggestion_message(
                    story_id=match.story.id,
                    title=match.story.title,
                    situation=match.story.situation,
                    relevance_score=match.relevance_score,
                    suggested_opening=match.suggested_opening,
                    key_metrics=match.key_metrics,
                    tags=match.story.tags
                )
                await self.broadcast(msg)
        except Exception as e:
            logger.warning(f"Story recall failed: {e}")

    async def _suggest_structure(self, question: str, q_type: str) -> None:
        """Suggest an answer structure framework."""
        if not self.structure_suggester:
            return
            
        try:
            hint = self.structure_suggester.suggest_structure(question, q_type)
            msg = create_structure_suggestion_message(
                name=hint.name,
                sections=[s.to_dict() for s in hint.sections],
                tips=hint.tips
            )
            await self.broadcast(msg)
        except Exception as e:
            logger.warning(f"Structure suggestion failed: {e}")

    async def _process_question_pipeline(
        self,
        original_question: str,
        question_type: str,
        start_time: float
    ) -> None:
        """
        Phase 3C: Enhanced question processing pipeline.
        
        Pipeline: Query Reformulation → Question Splitting → Enhanced Retrieval → Answer Generation
        
        Args:
            original_question: The raw question text from transcription
            question_type: The detected question type (behavioral, technical, etc.)
            start_time: Pipeline start time for latency tracking
        """
        import time
        
        # Refine question_type using granular classification from prompts.py
        # This ensures "interview_question" becomes "intro", "behavioral", "technical", etc.
        # which allows DOC_PRIORITY_BY_QUESTION_TYPE to properly prioritize SAMPLE_QA
        refined_question_type = classify_question(original_question)
        if refined_question_type != "general":
            logger.info(f"Refined question type: {question_type} → {refined_question_type}")
            question_type = refined_question_type
        
        # Phase 4E: Recall relevant stories (parallel)
        if self.story_recaller and question_type in ("behavioral", "interview_question", "intro", "weakness", "conflict", "leadership"):
            asyncio.create_task(self._recall_and_suggest_story(original_question, question_type))
            
        # Phase 4E: Suggest structure (parallel)
        asyncio.create_task(self._suggest_structure(original_question, question_type))
        
        # Step 1: Query Reformulation (expand follow-ups)
        reformulated_question = original_question
        try:
            # Build conversation history in format expected by reformulator
            history_for_reformulator = []
            for i in range(0, len(self.session_state.conversation_history) - 1, 2):
                if i + 1 < len(self.session_state.conversation_history):
                    history_for_reformulator.append({
                        "question": self.session_state.conversation_history[i].get("content", ""),
                        "answer": self.session_state.conversation_history[i + 1].get("content", "")
                    })
            
            reformulated_question, was_reformulated = self.query_reformulator.reformulate_if_needed(
                original_question, history_for_reformulator
            )
            if was_reformulated:
                logger.info(f"Reformulated: '{original_question[:50]}...' → '{reformulated_question[:50]}...'")
        except Exception as e:
            logger.warning(f"Query reformulation failed: {e}")
            reformulated_question = original_question
        
        # Step 2: Question Splitting (handle compound questions)
        sub_questions = [reformulated_question]
        try:
            split_result = self.question_splitter.split_questions(reformulated_question)
            if len(split_result) > 1:
                sub_questions = split_result
                logger.info(f"Split into {len(sub_questions)} sub-questions: {sub_questions}")
        except Exception as e:
            logger.warning(f"Question splitting failed: {e}")
        
        # Step 3: Retrieve context (using question type for enhanced retrieval when available)
        context_chunks, rag_confidence = await self._retrieve_context_enhanced(
            reformulated_question, question_type, sub_questions
        )
        
        # Step 4: Generate answer
        await self._generate_answer_with_context(
            original_question=original_question,
            reformulated_question=reformulated_question,
            context_chunks=context_chunks,
            rag_confidence=rag_confidence,
            start_time=start_time
        )
    
    async def _retrieve_context_enhanced(
        self,
        question: str,
        question_type: str,
        sub_questions: List[str]
    ) -> tuple[list[str], ConfidenceLevel]:
        """
        Enhanced context retrieval using question type, sub-questions, and speculative cache.
        """
        # Phase 4D: Use speculative results if available
        if self.speculative_retriever:
            try:
                results = await self.speculative_retriever.on_segment_complete(question)
                if results:
                    chunks = [r.text for r in results]
                    confidence = self._confidence_from_string(results[0].confidence)
                    logger.info(f"Using {len(chunks)} speculative/retrieved chunks")
                    return chunks, confidence
            except Exception as e:
                logger.warning(f"Speculative retrieval failed in pipeline: {e}")
        
        # Use EnhancedRAGEngine filtering if available
        if self.rag_engine and hasattr(self.rag_engine, 'retrieve_for_question'):
            try:
                results = self.rag_engine.retrieve_for_question(
                    question=question,
                    question_type=question_type,
                    sub_questions=sub_questions,
                    limit=5
                )
                if results:
                    chunks = [r.text for r in results]
                    confidence = self._confidence_from_string(results[0].confidence)
                    logger.info(f"Enhanced retrieval ({question_type}): {len(chunks)} chunks")
                    return chunks, confidence
            except Exception as e:
                logger.error(f"Enhanced retrieval failed: {e}")
        
        # Fallback to standard retrieval
        return self._retrieve_context(question)
    
    async def _check_consistency(self, text: str) -> None:
        """Extract claims and check for contradictions."""
        if not self.consistency_tracker or not self.session_state.persistent_session_id:
            return
            
        try:
            result = self.consistency_tracker.extract_and_check(text)
            if result.contradictions:
                msg = create_consistency_warning_message([c.to_dict() for c in result.contradictions])
                await self.broadcast(msg)
        except Exception as e:
            logger.warning(f"Consistency check failed: {e}")

    async def _generate_answer_with_context(
        self,
        original_question: str,
        reformulated_question: str,
        context_chunks: List[str],
        rag_confidence: ConfidenceLevel,
        start_time: float
    ) -> None:
        """
        Generate answer with provided context and persist.
        
        Args:
            original_question: Original transcribed question
            reformulated_question: Reformulated question (may be same as original)
            context_chunks: Retrieved context chunks
            rag_confidence: Confidence level from retrieval
            start_time: Pipeline start time for latency tracking
        """
        import time
        
        if not self.llm:
            logger.warning("LLM not initialized, cannot generate answer")
            return
        
        try:
            context_str = "\n\n".join(context_chunks)
            full_answer_parts: List[str] = []
            
            # Use reformulated question for generation but track original
            async for chunk in self.llm.generate_response(
                reformulated_question, context_str, self.session_state.conversation_history
            ):
                full_answer_parts.append(chunk)
                answer_msg = create_answer_chunk_message(chunk=chunk, complete=False)
                await self.broadcast(answer_msg)
            
            latency = time.time() - start_time
            logger.info(f"Answer generation complete. Latency: {latency:.2f}s")
            
            final_msg = create_answer_chunk_message(
                chunk="",
                complete=True,
                confidence=rag_confidence
            )
            await self.broadcast(final_msg)
            
            # Append Q&A to conversation history (use original question for history)
            full_answer = "".join(full_answer_parts)
            
            # Phase 4E: Check consistency
            self._create_background_task(self._check_consistency(full_answer))
            
            self.session_state.conversation_history.append({
                "role": "user",
                "content": original_question
            })
            self.session_state.conversation_history.append({
                "role": "assistant",
                "content": full_answer
            })
            logger.debug(f"Added Q&A to history. Total exchanges: {len(self.session_state.conversation_history) // 2}")
            
            # Phase 3: Persist answer to session history
            if self.session_persistence_enabled and self.session_state.persistent_session_id:
                try:
                    self.session_store.add_answer(
                        session_id=self.session_state.persistent_session_id,
                        question=original_question,
                        answer=full_answer,
                        confidence=rag_confidence.value,
                        rag_chunks=context_chunks[:3] if context_chunks else None,
                        latency_ms=int(latency * 1000)
                    )
                except Exception as e:
                    logger.warning(f"Failed to persist answer: {e}")
                    
        except Exception as e:
            logger.error(f"LLM generation failed: {e}")
            error_msg = create_error_message(
                f"Failed to generate answer: {e}",
                code="ERR_LLM_GENERATION"
            )
            await self.broadcast(error_msg)

    async def _generate_answer_for_question(self, question: str, start_time: float) -> None:
        """Generate answer using RAG retrieval + LLM streaming."""
        import time
        
        context_chunks, rag_confidence = self._retrieve_context(question)
        
        if not self.llm:
            logger.warning("LLM not initialized, cannot generate answer")
            return
            
        try:
            context_str = "\n\n".join(context_chunks)
            # Collect full answer for history
            full_answer_parts: List[str] = []
            
            async for chunk in self.llm.generate_response(
                question, context_str, self.session_state.conversation_history
            ):
                full_answer_parts.append(chunk)
                answer_msg = create_answer_chunk_message(chunk=chunk, complete=False)
                await self.broadcast(answer_msg)
            
            latency = time.time() - start_time
            logger.info(f"Answer generation complete. Latency: {latency:.2f}s")
            
            final_msg = create_answer_chunk_message(
                chunk="",
                complete=True,
                confidence=rag_confidence
            )
            await self.broadcast(final_msg)
            
            # Append Q&A to conversation history for future context
            full_answer = "".join(full_answer_parts)
            
            # Phase 4E: Check consistency
            self._create_background_task(self._check_consistency(full_answer))
            
            self.session_state.conversation_history.append({
                "role": "user",
                "content": question
            })
            self.session_state.conversation_history.append({
                "role": "assistant", 
                "content": full_answer
            })
            logger.debug(f"Added Q&A to history. Total exchanges: {len(self.session_state.conversation_history) // 2}")
            
            # Phase 3: Persist answer to session history
            if self.session_persistence_enabled and self.session_state.persistent_session_id:
                try:
                    self.session_store.add_answer(
                        session_id=self.session_state.persistent_session_id,
                        question=question,
                        answer=full_answer,
                        confidence=rag_confidence.value,
                        rag_chunks=context_chunks[:3] if context_chunks else None,  # Top 3 chunks
                        latency_ms=int(latency * 1000)
                    )
                except Exception as e:
                    logger.warning(f"Failed to persist answer: {e}")
            
        except Exception as e:
            logger.error(f"LLM generation failed: {e}")
            error_msg = create_error_message(
                f"Failed to generate answer: {e}",
                code="ERR_LLM_GENERATION"
            )
            await self.broadcast(error_msg)

    # ===================
    # Phase 8: RAG Persistence Handlers
    # ===================
    
    async def _handle_load_rag_state(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """
        Handle LOAD_RAG_STATE message.
        
        Returns the current state of persistent RAG data, enabling the frontend
        to display existing documents without re-uploading on app restart.
        """
        try:
            state_summary = self.rag_manifest.get_state_summary()
            
            # Prepare document info for frontend
            documents = []
            for doc in state_summary.get("documents", []):
                documents.append({
                    "filename": doc.get("filename", ""),
                    "documentType": doc.get("document_type", "custom"),
                    "uploadTimestamp": doc.get("upload_timestamp", ""),
                    "fileSizeBytes": doc.get("file_size_bytes", 0),
                    "chunkCount": doc.get("chunk_count", 0),
                    "preview": doc.get("preview", ""),
                })
            
            response_msg = create_rag_state_message(
                has_documents=state_summary.get("has_documents", False),
                document_count=state_summary.get("document_count", 0),
                documents=documents,
                cache_expired=state_summary.get("cache_expired", True),
                last_cache_timestamp=state_summary.get("last_cache_timestamp"),
            )
            await websocket.send(response_msg.to_json())
            
            logger.info(f"RAG state loaded: {state_summary.get('document_count', 0)} documents, "
                       f"cache_expired={state_summary.get('cache_expired', True)}")
            
        except Exception as e:
            logger.error(f"Failed to load RAG state: {e}")
            error_msg = create_error_message(
                f"Failed to load RAG state: {e}",
                code="ERR_RAG_STATE"
            )
            await websocket.send(error_msg.to_json())
    
    async def _handle_refresh_cache(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """
        Handle REFRESH_CACHE message.
        
        Re-creates Gemini cache from existing documents without requiring
        the user to re-upload files. Uses stored base64 content from manifest.
        """
        try:
            data = message.data or {}
            
            # Get API keys from message if provided (for provider initialization)
            if "apiKeys" in data and not self.gemini_cache_manager:
                config = ProviderConfig.from_dict(data)
                if config.gemini_api_key:
                    self.gemini_cache_manager = GeminiCacheManager(config.gemini_api_key)
                    self.gemini_file_uploader = GeminiFileUploader(config.gemini_api_key)
                    logger.info("Initialized Gemini managers for cache refresh")
            
            if not self.gemini_cache_manager or not self.gemini_file_uploader:
                raise RuntimeError("Gemini managers not initialized. Provide API key.")
            
            # Get documents with stored content
            docs_for_cache = self.rag_manifest.get_documents_for_cache()
            
            if not docs_for_cache:
                raise RuntimeError("No documents with stored content available for cache refresh. "
                                  "Please re-upload documents.")
            
            logger.info(f"Refreshing cache with {len(docs_for_cache)} documents...")
            
            # Re-upload files to Gemini
            file_doc_type_map = {
                "resume": FileDocumentType.RESUME,
                "job_description": FileDocumentType.JOB_DESCRIPTION,
                "company_info": FileDocumentType.COMPANY_INFO,
                "interviewer_info": FileDocumentType.INTERVIEWER_INFO,
                "sample_qa": FileDocumentType.SAMPLE_QA,
                "industry_research": FileDocumentType.INDUSTRY_RESEARCH,
                "custom": FileDocumentType.CUSTOM,
            }
            
            # Clear existing uploads before refreshing
            self.gemini_file_uploader.clear()
            
            # Upload all documents in parallel
            upload_coros = []
            for doc in docs_for_cache:
                if doc.get("content_b64"):
                    file_doc_type = file_doc_type_map.get(
                        doc.get("document_type", "custom"), 
                        FileDocumentType.CUSTOM
                    )
                    upload_coros.append(
                        self.gemini_file_uploader.upload_from_base64_async(
                            content_b64=doc["content_b64"],
                            filename=doc["filename"],
                            document_type=file_doc_type
                        )
                    )
            
            if upload_coros:
                await asyncio.gather(*upload_coros, return_exceptions=True)
                logger.info(f"Re-uploaded {len(upload_coros)} files to Gemini")
            
            # Create new cache
            if self.gemini_file_uploader.has_files():
                uploaded_files = self.gemini_file_uploader.get_uploaded_files()
                document_manifest = self.gemini_file_uploader.get_document_manifest()
                
                # Get profile for cache if available
                profile_text = None
                if self.memory_store:
                    profile = self.memory_store.get_profile()
                    if profile:
                        profile_text = profile.get_prompt_injection()
                
                cache_name = await self.gemini_cache_manager.create_cache_from_files_async(
                    uploaded_files=uploaded_files,
                    document_manifest=document_manifest,
                    ttl_seconds=7200,  # 2 hours
                    model=GeminiModels.DEFAULT_LLM,
                    profile_text=profile_text,
                )
                
                # Update manifest with new cache timestamp
                self.rag_manifest.update_cache_timestamp()
                
                # Attach cache to LLM if available
                if cache_name and self.llm and hasattr(self.llm, 'set_cached_content'):
                    getattr(self.llm, 'set_cached_content')(cache_name)
                    logger.info(f"Attached refreshed cache to LLM: {cache_name}")
                
                response_msg = create_cache_refresh_complete_message(
                    success=True,
                    cache_name=cache_name
                )
                await websocket.send(response_msg.to_json())
                
                logger.info(f"Cache refresh complete: {cache_name}")
            else:
                raise RuntimeError("No files uploaded to Gemini after refresh attempt")
                
        except Exception as e:
            logger.error(f"Cache refresh failed: {e}")
            response_msg = create_cache_refresh_complete_message(
                success=False,
                error=str(e)
            )
            await websocket.send(response_msg.to_json())
    
    async def _handle_clear_all_data(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """
        Handle CLEAR_ALL_DATA message.
        
        Wipes all persistent data:
        - ChromaDB collection (RAG embeddings)
        - MemoryStore (profile, stories, facts)
        - RAG manifest
        - Gemini file uploads (if active)
        - Gemini cache (if active)
        """
        try:
            cleared_items = {
                "chroma_collection": False,
                "memory_store": False,
                "rag_manifest": False,
                "gemini_files": False,
                "gemini_cache": False,
            }
            
            # Clear ChromaDB collection
            if self.vector_store:
                try:
                    self.vector_store.clear()
                    cleared_items["chroma_collection"] = True
                    logger.info("Cleared ChromaDB collection")
                except Exception as e:
                    logger.warning(f"Failed to clear ChromaDB: {e}")
            
            # Clear Memory Store
            if self.memory_store:
                try:
                    self.memory_store.clear_all()
                    cleared_items["memory_store"] = True
                    logger.info("Cleared Memory Store")
                except Exception as e:
                    logger.warning(f"Failed to clear Memory Store: {e}")
            
            # Clear RAG Manifest
            try:
                self.rag_manifest.clear()
                cleared_items["rag_manifest"] = True
                logger.info("Cleared RAG Manifest")
            except Exception as e:
                logger.warning(f"Failed to clear RAG Manifest: {e}")
            
            # Clear Gemini File Uploader
            if self.gemini_file_uploader:
                try:
                    self.gemini_file_uploader.clear()
                    cleared_items["gemini_files"] = True
                    logger.info("Cleared Gemini file uploads")
                except Exception as e:
                    logger.warning(f"Failed to clear Gemini files: {e}")
            
            # Clear Gemini Cache
            if self.gemini_cache_manager:
                try:
                    self.gemini_cache_manager.delete_current_cache()
                    cleared_items["gemini_cache"] = True
                    logger.info("Cleared Gemini cache")
                except Exception as e:
                    logger.warning(f"Failed to clear Gemini cache: {e}")
            
            # Clear Context Manager
            if self.context_manager:
                try:
                    self.context_manager.clear_context()
                    logger.info("Cleared Context Manager")
                except Exception as e:
                    logger.warning(f"Failed to clear Context Manager: {e}")
            
            # Clear LLM candidate profile
            if self.llm and hasattr(self.llm, 'clear_candidate_profile'):
                try:
                    self.llm.clear_candidate_profile()
                    logger.info("Cleared LLM candidate profile")
                except Exception as e:
                    logger.warning(f"Failed to clear LLM profile: {e}")
            
            response_msg = create_data_cleared_message(
                success=True,
                cleared_items=cleared_items
            )
            await websocket.send(response_msg.to_json())
            
            logger.info(f"All data cleared: {cleared_items}")
            
        except Exception as e:
            logger.error(f"Failed to clear all data: {e}")
            response_msg = create_data_cleared_message(
                success=False,
                error=str(e)
            )
            await websocket.send(response_msg.to_json())

    async def _handle_pause_listening(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """
        Handle PAUSE_LISTENING message.
        
        Stops audio capture and streaming STT while keeping the session active.
        Manual question posting, RAG, and LLM will continue to work.
        """
        if self.session_state.listening_paused:
            logger.info("Listening already paused")
            return
            
        try:
            # Stop audio processing
            await self._stop_audio_processing()
            logger.info("Audio processing stopped")
            
            # Stop streaming STT
            if self.streaming_stt_manager:
                try:
                    await self.streaming_stt_manager.stop_session()
                    logger.info("Stopped streaming STT session")
                except Exception as e:
                    logger.warning(f"Error stopping streaming STT: {e}")
            
            # Update state
            self.session_state.listening_paused = True
            self.session_state.status = SessionStatus.LISTENING_PAUSED
            
            # Notify client
            response_msg = Message(
                type=MessageType.LISTENING_PAUSED,
                data={"success": True}
            )
            await self.broadcast(response_msg)
            
            # Also send status update
            status_msg = create_status_message(SessionStatus.LISTENING_PAUSED)
            await self.broadcast(status_msg)
            
            logger.info("Listening paused - manual input only mode")
            
        except Exception as e:
            logger.error(f"Failed to pause listening: {e}")
            error_msg = create_error_message(
                f"Failed to pause listening: {e}",
                code="ERR_PAUSE_FAILED"
            )
            await websocket.send(error_msg.to_json())
    
    async def _handle_resume_listening(
        self,
        websocket: ServerConnection,
        message: Message
    ) -> None:
        """
        Handle RESUME_LISTENING message.
        
        Restarts audio capture and streaming STT.
        """
        if not self.session_state.listening_paused:
            logger.info("Listening not paused, nothing to resume")
            return
            
        try:
            # Restart audio processing
            await self._start_audio_processing()
            logger.info("Audio processing restarted")
            
            # Restart streaming STT if enabled
            if self.streaming_stt_enabled and self.provider_factory:
                try:
                    if not self.streaming_stt_manager:
                        self.streaming_stt_manager = StreamingSTTManager(self.provider_factory)
                    
                    callbacks = StreamingSTTCallbacks(
                        on_interim=self._on_streaming_interim,
                        on_final=self._on_streaming_final,
                        on_end_of_turn=self._on_streaming_end_of_turn,
                        on_error=self._on_streaming_error,
                    )
                    
                    streaming_started = await self.streaming_stt_manager.start_session(callbacks)
                    if streaming_started:
                        logger.info(f"Streaming STT restarted: {self.streaming_stt_manager.provider_name}")
                    else:
                        logger.info("Streaming STT not available, using batch STT only")
                except Exception as e:
                    logger.warning(f"Failed to restart streaming STT: {e}")
            
            # Update state
            self.session_state.listening_paused = False
            self.session_state.status = SessionStatus.LISTENING
            
            # Notify client
            response_msg = Message(
                type=MessageType.LISTENING_RESUMED,
                data={"success": True}
            )
            await self.broadcast(response_msg)
            
            # Also send status update
            status_msg = create_status_message(SessionStatus.LISTENING)
            await self.broadcast(status_msg)
            
            logger.info("Listening resumed - live transcription active")
            
        except Exception as e:
            logger.error(f"Failed to resume listening: {e}")
            error_msg = create_error_message(
                f"Failed to resume listening: {e}",
                code="ERR_RESUME_FAILED"
            )
            await websocket.send(error_msg.to_json())

    def _retrieve_context(self, question: str) -> tuple[list[str], ConfidenceLevel]:
        """
        Retrieve context chunks from RAG engine for the given question.
        
        Cache-First Architecture:
        - If Gemini cache is available (file-based), skip RAG entirely
        - Cache contains full documents with proper attribution
        - RAG would be redundant and could cause context confusion
        """
        context_chunks = []
        rag_confidence = ConfidenceLevel.LOW
        
        # Cache-First: Skip RAG if LLM has cached content
        # The Gemini cache contains full documents, RAG chunks are redundant
        if self.llm and hasattr(self.llm, '_cached_content_name'):
            cached_content = getattr(self.llm, '_cached_content_name', None)
            if cached_content:
                logger.info("Cache-First: Skipping RAG retrieval (using Gemini file cache)")
                return [], ConfidenceLevel.HIGH  # High confidence because cache has full docs
        
        if not self.rag_engine:
            return context_chunks, rag_confidence
            
        try:
            retrieval_results = self.rag_engine.retrieve(question, limit=5)
            logger.info(f"Retrieved {len(retrieval_results)} chunks for question")
            
            if retrieval_results:
                context_chunks = [r.text for r in retrieval_results]
                rag_confidence = self._confidence_from_string(retrieval_results[0].confidence)
                
                for i, r in enumerate(retrieval_results):
                    logger.debug(f"Chunk {i}: {r.confidence} ({r.distance:.2f}) - {r.text[:50]}...")
        except Exception as e:
            logger.error(f"RAG retrieval failed: {e}")
            
        return context_chunks, rag_confidence

    def _confidence_from_string(self, confidence_str: str) -> ConfidenceLevel:
        """Convert string confidence to ConfidenceLevel enum."""
        mapping = {
            "high": ConfidenceLevel.HIGH,
            "medium": ConfidenceLevel.MEDIUM,
            "low": ConfidenceLevel.LOW,
        }
        return mapping.get(confidence_str.lower(), ConfidenceLevel.LOW)


async def main() -> None:
    """Main entry point for the sidecar server."""
    server = SidecarServer()

    try:
        await server.start()
    except KeyboardInterrupt:
        logger.info("Received shutdown signal")
    finally:
        await server.stop()


if __name__ == "__main__":
    asyncio.run(main())
